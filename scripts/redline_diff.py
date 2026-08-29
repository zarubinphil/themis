#!/usr/bin/env python3
"""redline_diff.py — разбор правок доверителя в .docx ПО СТРУКТУРЕ, не по байтам.

Зачем. Доверитель правит выданный документ под себя, и Фемида обязана из этих
правок учиться (`knowledge/redlines.md`). Но `.docx` — это zip: байтовое сравнение
двух файлов всегда кричит «отличается» (пересохранение меняет метаданные и порядок
XML) и никогда не говорит, ЧТО именно изменилось. Учиться на «файлы разные»
нельзя. Этот прибор вскрывает оба документа через python-docx и отвечает предметно:
какие абзацы переписаны/добавлены/убраны (смысл) и что стало со шрифтом, полями,
выравниванием (форма) — ровно две оси, по которым `redlines.md` копит уроки.

Почему не diff по XML. Сырой XML документа шумит на служебных атрибутах (rsid,
порядок свойств), которых доверитель не касался. Идем через объектную модель:
берем то, что человек реально менял — текст абзацев, гарнитуру, поля, выравнивание.
"""
from __future__ import annotations

import argparse
import json
import sys
from difflib import SequenceMatcher

from docx import Document

# Значение выравнивания python-docx → слово. None — свойство не задано (наследуется).
ALIGN = {0: "слева", 1: "центр", 2: "справа", 3: "по ширине", None: "по умолчанию"}


def _align(p) -> str:
    a = p.alignment
    return ALIGN.get(int(a) if a is not None else None, str(a))


def extract(path: str) -> dict:
    """Структура документа: абзацы (текст+выравнивание), поля, набор гарнитур."""
    doc = Document(path)
    paras = [{"text": p.text.strip(), "align": _align(p)}
             for p in doc.paragraphs if p.text.strip()]

    sec = doc.sections[0]
    margins = {}
    for name, val in (("верх", sec.top_margin), ("низ", sec.bottom_margin),
                      ("лево", sec.left_margin), ("право", sec.right_margin)):
        if val is not None:
            margins[name] = round(val.mm, 1)

    fonts = set()
    normal = doc.styles["Normal"].font.name if "Normal" in doc.styles else None
    if normal:
        fonts.add(normal)
    for p in doc.paragraphs:
        for r in p.runs:
            if r.font.name:
                fonts.add(r.font.name)
    return {"paras": paras, "margins": margins, "fonts": sorted(fonts)}


def content_diff(before: list[dict], after: list[dict]) -> list[dict]:
    """Что изменилось ПО СМЫСЛУ: переписанные, добавленные, убранные абзацы."""
    a = [p["text"] for p in before]
    b = [p["text"] for p in after]
    out = []
    for tag, i1, i2, j1, j2 in SequenceMatcher(None, a, b).get_opcodes():
        if tag == "equal":
            continue
        if tag == "replace":
            old, new = a[i1:i2], b[j1:j2]
            for k in range(max(len(old), len(new))):
                out.append({
                    "op": "переписан",
                    "было": old[k] if k < len(old) else "",
                    "стало": new[k] if k < len(new) else "",
                })
        elif tag == "delete":
            out += [{"op": "убран", "было": t, "стало": ""} for t in a[i1:i2]]
        elif tag == "insert":
            out += [{"op": "добавлен", "было": "", "стало": t} for t in b[j1:j2]]
    return out


def format_diff(before: dict, after: dict) -> list[dict]:
    """Что изменилось ПО ФОРМЕ: поля, гарнитура, выравнивание совпавших абзацев."""
    out = []
    for name in ("верх", "низ", "лево", "право"):
        b, a = before["margins"].get(name), after["margins"].get(name)
        if b != a:
            out.append({"тип": "поле", "поле": name, "было": b, "стало": a})

    if before["fonts"] != after["fonts"]:
        out.append({"тип": "шрифт", "было": before["fonts"], "стало": after["fonts"]})

    # Выравнивание сравниваем только у абзацев, чей ТЕКСТ не менялся: иначе «другое
    # выравнивание» неотличимо от «другой абзац» и учить на нем нечего.
    a_txt = [p["text"] for p in before["paras"]]
    b_txt = [p["text"] for p in after["paras"]]
    for tag, i1, i2, j1, j2 in SequenceMatcher(None, a_txt, b_txt).get_opcodes():
        if tag != "equal":
            continue
        for off in range(i2 - i1):
            bp, ap = before["paras"][i1 + off], after["paras"][j1 + off]
            if bp["align"] != ap["align"]:
                out.append({"тип": "выравнивание", "абзац": bp["text"][:40],
                            "было": bp["align"], "стало": ap["align"]})
    return out


def diff(before_path: str, after_path: str) -> dict:
    b, a = extract(before_path), extract(after_path)
    return {"content": content_diff(b["paras"], a["paras"]),
            "format": format_diff(b, a)}


def _build_sample(tmp: str) -> tuple[str, str]:
    """Два docx с известными правками — смысловыми и форматными — без сети."""
    import os

    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Mm

    d1 = Document()
    d1.sections[0].left_margin = Mm(30)
    d1.styles["Normal"].font.name = "Times New Roman"
    p = d1.add_paragraph("Первый абзац без изменений.")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    d1.add_paragraph("Этот абзац доверитель удалит.")
    d1.add_paragraph("Формулировка старая.")
    before = os.path.join(tmp, "before.docx")
    d1.save(before)

    d2 = Document()
    d2.sections[0].left_margin = Mm(35)          # поле изменено 30 → 35
    d2.styles["Normal"].font.name = "PT Serif"   # гарнитура изменена
    p = d2.add_paragraph("Первый абзац без изменений.")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER      # выравнивание изменено
    d2.add_paragraph("Формулировка новая, доверитель переписал.")
    after = os.path.join(tmp, "after.docx")
    d2.save(after)
    return before, after


def selftest() -> int:
    import tempfile
    with tempfile.TemporaryDirectory() as tmp:
        before, after = _build_sample(tmp)
        r = diff(before, after)
    c, f = r["content"], r["format"]
    content_texts = " ".join(e["было"] + e["стало"] for e in c)
    fmt_types = {e["тип"] for e in f}
    # Чистое удаление (соседних правок нет) обязано дать метку «убран» — проверяем
    # на голых списках, без docx: в общем docx-примере SequenceMatcher сливает
    # удаление в соседний replace, и это верно, но метку «убран» так не увидеть.
    pure_del = content_diff([{"text": "A"}, {"text": "B"}, {"text": "C"}],
                            [{"text": "A"}, {"text": "C"}])
    checks = [
        ("правки по смыслу пойманы", len(c) > 0),
        ("удаленный текст виден в разборе", any("удалит" in e["было"] for e in c)),
        ("чистое удаление помечается «убран»",
         any(e["op"] == "убран" and e["было"] == "B" for e in pure_del)),
        ("переписанный абзац виден", "Формулировка новая" in content_texts),
        ("неизменный абзац в смысловой diff не попал", "без изменений" not in content_texts),
        ("изменение поля поймано", any(e.get("поле") == "лево" and e["было"] == 30.0
                                       and e["стало"] == 35.0 for e in f)),
        ("смена гарнитуры поймана", "шрифт" in fmt_types),
        ("смена выравнивания у совпавшего абзаца поймана",
         any(e["тип"] == "выравнивание" and e["было"] == "по ширине"
             and e["стало"] == "центр" for e in f)),
    ]
    bad = [n for n, ok in checks if not ok]
    for n, ok in checks:
        print(f"  {'✓' if ok else '✗'} {n}")
    if bad:
        print(f"selftest ПРОВАЛЕН: {len(bad)} из {len(checks)}")
        return 1
    print(f"selftest пройден: {len(checks)}/{len(checks)} — структурный diff без сети")
    return 0


def main() -> int:
    ap = argparse.ArgumentParser(description="Разбор правок доверителя в .docx по структуре")
    ap.add_argument("before", nargs="?", metavar="ДО.docx")
    ap.add_argument("after", nargs="?", metavar="ПОСЛЕ.docx")
    ap.add_argument("--json", action="store_true", help="машинный вывод {content, format}")
    ap.add_argument("--selftest", action="store_true", help="проверка без сети")
    a = ap.parse_args()

    if a.selftest:
        return selftest()
    if not a.before or not a.after:
        ap.error("нужны ДО.docx и ПОСЛЕ.docx (или --selftest)")
    result = diff(a.before, a.after)
    print(json.dumps(result, ensure_ascii=False, indent=2))
    return 0


if __name__ == "__main__":
    sys.exit(main())
