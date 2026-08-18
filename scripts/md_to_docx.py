#!/usr/bin/env python3
"""
Универсальный конвертер .md -> .docx через DocBuilder.
Использование:
    python3 scripts/md_to_docx.py cases/.../.agent/drafts/document.md
Сохраняет document.docx рядом с исходным .md.
"""

import re
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent.parent))
from scripts.create_docx import DocBuilder, FONT, _set_font
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def strip_markdown_bold(text: str) -> list:
    """Превращает строку с **жирным** текстом в список (text, bold)."""
    parts = []
    while "**" in text:
        idx = text.index("**")
        if idx > 0:
            parts.append((text[:idx], False))
        text = text[idx + 2:]
        if "**" in text:
            idx = text.index("**")
            parts.append((text[:idx], True))
            text = text[idx + 2:]
        else:
            parts.append((text, True))
            text = ""
    if text:
        parts.append((text, False))
    return parts


DOCUMENT_TITLES = {
    "ВОЗРАЖЕНИЯ", "ВОЗРАЖЕНИЕ", "ЗАЯВЛЕНИЕ", "ИСКОВОЕ ЗАЯВЛЕНИЕ",
    "ХОДАТАЙСТВО", "ОПРЕДЕЛЕНИЕ", "ПРЕДЛОЖЕНИЕ",
}


def is_document_heading(line: str) -> bool:
    """Определяет, заканчивается ли шапка на данной строке."""
    stripped = line.strip()
    if re.match(r"^#{1,4} \S", stripped):
        return True
    if stripped in DOCUMENT_TITLES:
        return True
    if re.match(r"^[IVXLC]+\.\s", stripped):
        return True
    return False


def parse_header(lines: list) -> dict:
    """Парсит шапку документа (до первого ---, заголовка или раздела)."""
    court_name = ""
    court_route_parts = []
    case_number = ""
    parties = []
    current_party = None

    for line in lines:
        line = line.rstrip()
        if not line:
            if current_party is not None:
                parties.append(current_party)
                current_party = None
            continue

        # суд
        if line.startswith("В ") and "суд" in line.lower():
            court_name = line[2:].strip()
            continue

        # адрес / судья
        if line.lower().startswith("адрес:") or line.lower().startswith("судья"):
            court_route_parts.append(line)
            continue

        # дело
        if line.lower().startswith("дело №"):
            case_number = line
            continue

        # метка стороны (заканчивается двоеточием)
        if line.endswith(":") and ("истец" in line.lower() or "ответчик" in line.lower() or "третье" in line.lower() or "лицо" in line.lower()):
            if current_party is not None:
                parties.append(current_party)
            current_party = {"label": line, "lines": []}
            continue

        # строка данных стороны
        if current_party is not None:
            clean = line.rstrip(",;").strip()
            # первая строка блока — имя, делаем жирным
            is_bold = len(current_party["lines"]) == 0
            current_party["lines"].append((clean, is_bold))

    if current_party is not None:
        parties.append(current_party)

    court_route = "\n".join(court_route_parts)
    return {
        "court": court_name,
        "court_route": court_route,
        "case": case_number,
        "parties": parties,
    }


def add_table_to_doc(doc, rows: list):
    """Добавляет markdown-таблицу в docx с базовым форматированием."""
    if not rows:
        return
    # отфильтровать разделительные строки
    data_rows = [r for r in rows if not re.match(r"^\|[-:\|\s]+\|$", r.strip())]
    if not data_rows:
        return

    cells = [list(map(str.strip, row.strip("|").split("|"))) for row in data_rows]
    num_cols = max(len(c) for c in cells)
    table = doc.add_table(rows=len(cells), cols=num_cols)

    # убрать границы
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement("w:tblBorders")
    for side in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "none")
        tblBorders.append(el)
    tblPr.append(tblBorders)

    # автоширина
    tblLayout = OxmlElement("w:tblLayout")
    tblLayout.set(qn("w:type"), "autofit")
    tblPr.append(tblLayout)

    for i, row_cells in enumerate(cells):
        for j in range(num_cols):
            cell = table.rows[i].cells[j]
            cell.paragraphs[0].clear()
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
            text = row_cells[j] if j < len(row_cells) else ""
            _set_font(cell.paragraphs[0].add_run(text), 12)

    # отступ после таблицы
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)


def convert(md_path: str):
    md_file = Path(md_path)
    out_path = md_file.with_suffix(".docx")

    content = md_file.read_text(encoding="utf-8")
    lines = content.splitlines()

    # разделить шапку и тело
    header_lines = []
    body_lines = []
    in_header = True
    for line in lines:
        if in_header and line.strip() == "---":
            in_header = False
            continue
        if in_header and is_document_heading(line):
            in_header = False
            body_lines.append(line)
            continue
        if in_header:
            header_lines.append(line)
        else:
            body_lines.append(line)

    header_info = parse_header(header_lines)

    b = DocBuilder()

    # шапка
    b.add_header_table(
        court_name=header_info["court"] or "Суд",
        court_route=header_info["court_route"],
        parties=header_info["parties"],
        case_number=header_info["case"],
    )
    b.add_empty()

    # тело
    i = 0
    state = "body"  # body | proshyu | appendices

    while i < len(body_lines):
        line = body_lines[i]
        stripped = line.strip()

        # пустая строка
        if not stripped:
            i += 1
            continue

        # горизонтальная черта
        if stripped == "---":
            b.add_empty()
            i += 1
            continue

        # заголовок документа (# )
        if re.match(r"^# ", line):
            b.add_title(stripped[2:])
            i += 1
            continue

        # документный заголовок без решетки (ВОЗРАЖЕНИЯ, ЗАЯВЛЕНИЕ и т.д.)
        if stripped in DOCUMENT_TITLES:
            b.add_title(stripped)
            i += 1
            # следующие непустые строки до разделителя — подзаголовки
            while i < len(body_lines) and body_lines[i].strip() and not is_document_heading(body_lines[i]):
                b.add_subtitle(body_lines[i].strip())
                i += 1
            continue

        # подзаголовок (## )
        if re.match(r"^## ", line):
            b.add_subtitle(stripped[3:])
            i += 1
            continue

        # секция (### )
        if re.match(r"^### ", line):
            b.add_section(stripped[4:])
            i += 1
            continue

        # подсекция (#### )
        if re.match(r"^#### ", line):
            b.add_subsection(stripped[5:])
            i += 1
            continue

        # таблица
        if stripped.startswith("|"):
            table_rows = []
            while i < len(body_lines) and body_lines[i].strip().startswith("|"):
                table_rows.append(body_lines[i])
                i += 1
            add_table_to_doc(b.doc, table_rows)
            continue

        # маркированный список
        if stripped.startswith("- "):
            item = stripped[2:]
            b.add_bullet(item)
            i += 1
            continue

        # ПРОСИМ / ПРОШУ
        if stripped in ("ПРОСИМ:", "ПРОШУ:"):
            b.add_proshyu()
            state = "proshyu"
            i += 1
            continue

        # ПРИЛОЖЕНИЯ
        if stripped == "ПРИЛОЖЕНИЯ:":
            b.add_appendices()
            state = "appendices"
            i += 1
            continue

        # просительная часть: пункты вида "1. ..."
        if state == "proshyu" and re.match(r"^\d+\.\s", stripped):
            b.add_request_item(stripped)
            i += 1
            continue

        # приложения: пункты вида "1. ..."
        if state == "appendices" and re.match(r"^\d+\.\s", stripped):
            b.add_appendix_item(stripped)
            i += 1
            continue

        # подпись: строка с ролью и длинным подчеркиванием
        if "____________" in stripped or "________________" in stripped:
            # роль — предыдущая непустая строка, если есть
            role = ""
            name = stripped
            if i > 0:
                prev = body_lines[i - 1].strip()
                if prev and not prev.startswith("#") and prev not in ("ПРОСИМ:", "ПРОШУ:", "ПРИЛОЖЕНИЯ:"):
                    role = prev
            # дата — следующая непустая строка
            date_str = ""
            j = i + 1
            while j < len(body_lines) and not body_lines[j].strip():
                j += 1
            if j < len(body_lines):
                date_str = body_lines[j].strip()
            b.add_signature_table(role=role, name=name, date=date_str)
            i = j + 1 if date_str else i + 1
            continue

        # обычный абзац
        b.add_body(strip_markdown_bold(stripped))
        i += 1

    b.save(str(out_path))
    print(f"✓ Создано: {out_path}")


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Использование: python3 scripts/md_to_docx.py <путь к .md>")
        sys.exit(1)
    convert(sys.argv[1])
