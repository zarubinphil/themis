#!/usr/bin/env python3
"""
Создание судебного документа по эталону форматирования шаблон.docx.
Используется из doc-drafter. Все параметры верифицированы по XML шаблона.

Использование:
    from scripts.create_docx import DocBuilder
    b = DocBuilder()
    b.add_title("ИСКОВОЕ ЗАЯВЛЕНИЕ")
    b.add_subtitle("о взыскании неосновательного обогащения")
    b.add_section("I. ОБСТОЯТЕЛЬСТВА ДЕЛА")
    b.add_body("Текст абзаца...")
    b.add_proshyu()
    b.add_request_item("1. Взыскать с ответчика...")
    b.add_appendices()
    b.add_appendix_item("1. Договор от 01.01.2024")
    b.add_signature("Иванов Иван Иванович", "27.05.2026")
    b.save("путь/к/файлу.docx")
"""

import os
import sys

from docx import Document
from docx.shared import Pt, Cm, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ОДНА гарнитура на весь документ — PT Serif. Решение владельца 04.08.2026,
# отменяет четырехгарнитурный набор от 03.08.2026 (Playfair Display + PT Serif +
# Golos Text + PT Mono). PT Serif — свободный (SIL OFL), кириллица проектировалась
# первой; выбор закрывает требование ГОСТ Р 7.0.97-2016 п. 3.3 «необходимо
# использовать свободно распространяемые бесплатные шрифты» (Times New Roman —
# проприетарный Monotype).
#
# Имена констант оставлены: их читают полсотни мест кода, а роль — «титул»,
# «заголовок», «числовая колонка» — от смены гарнитуры никуда не делась.
# Иерархия теперь держится кеглем, начертанием и разрядкой, а не сменой шрифта.
# Менять только новым решением владельца: документы практики обязаны быть
# неотличимы друг от друга по оформлению.
FONT_BODY = "PT Serif"
FONT_DISPLAY = FONT_BODY
FONT_MONO = FONT_BODY
FONT_TITLE = FONT_BODY

# Разрядка титульного блока в пунктах (межбуквенный интервал).
TRACK_TITLE_PT = 3.0
TRACK_SUBTITLE_PT = 0.6

FONT = FONT_BODY  # обратная совместимость

# Порядок дочерних элементов w:rPr по ECMA-376. Элемент, вставленный не на
# свое место, схема считает недопустимым, и рендерер молча его игнорирует.
_RPR_ORDER = [
    "rStyle", "rFonts", "b", "bCs", "i", "iCs", "caps", "smallCaps", "strike",
    "dstrike", "outline", "shadow", "emboss", "imprint", "noProof",
    "snapToGrid", "vanish", "webHidden", "color", "spacing", "w", "kern",
    "position", "sz", "szCs", "highlight", "u", "effect", "bdr", "shd",
    "fitText", "vertAlign", "rtl", "cs", "em", "lang", "eastAsianLayout",
    "specVanish", "oMath",
]


def _rpr_insert(rPr, element):
    tag = element.tag.split("}")[-1]
    idx = _RPR_ORDER.index(tag)
    for child in rPr:
        child_tag = child.tag.split("}")[-1]
        if child_tag not in _RPR_ORDER or _RPR_ORDER.index(child_tag) > idx:
            child.addprevious(element)
            return
    rPr.append(element)


def _track(paragraph, pt):
    """Разрядка абзаца: w:spacing в rPr, в двадцатых долях пункта.

    QuickLook на macOS этот атрибут НЕ отрисовывает — проверять только в Word
    либо по XML (замер 03.08.2026). Отсутствие разрядки в предпросмотре
    ничего не доказывает.
    """
    for run in paragraph.runs:
        rPr = run._r.get_or_add_rPr()
        sp = OxmlElement("w:spacing")
        sp.set(qn("w:val"), str(int(round(pt * 20))))
        _rpr_insert(rPr, sp)
    return paragraph


def _set_font(run, size_pt, bold=False, family=None):
    run.font.name = family or FONT_BODY
    run.font.size = Pt(size_pt)
    run.font.bold = bold


def _set_docdefaults_line_spacing(doc):
    """Устанавливает межстрочный интервал 1.15× в docDefaults (как в шаблоне)."""
    settings = doc.settings.element
    # Найти или создать docDefaults
    docDefaults = settings.find(qn("w:docDefaults"))
    if docDefaults is None:
        # docDefaults должен быть в styles, не settings
        pass

    styles_elem = doc.styles.element
    docDefaults = styles_elem.find(qn("w:docDefaults"))
    if docDefaults is None:
        docDefaults = OxmlElement("w:docDefaults")
        styles_elem.insert(0, docDefaults)

    pPrDefault = docDefaults.find(qn("w:pPrDefault"))
    if pPrDefault is None:
        pPrDefault = OxmlElement("w:pPrDefault")
        docDefaults.append(pPrDefault)

    pPr = pPrDefault.find(qn("w:pPr"))
    if pPr is None:
        pPr = OxmlElement("w:pPr")
        pPrDefault.append(pPr)

    spacing = pPr.find(qn("w:spacing"))
    if spacing is None:
        spacing = OxmlElement("w:spacing")
        pPr.append(spacing)

    spacing.set(qn("w:line"), "276")
    spacing.set(qn("w:lineRule"), "auto")



# --------------------------------------------------------- нумерация Word
# Номера пунктов до 03.08.2026 писались обычным текстом («1. », «2. »), и
# Word их не пересчитывал: удалил пункт 3 — остальные приходилось
# перенумеровывать руками. Здесь заводится НАСТОЯЩАЯ нумерация Word
# (w:numPr + собственные определения в numbering.xml). Удаление, вставка и
# перестановка пункта пересчитываются программой сами.
#
# Каждый список получает свое определение: просительная часть, приложения и
# сквозные абзацы не должны делить счетчик.
# Идентификаторы продолжают нумерацию шаблона (10, 11, 12...), а НЕ уходят в
# отдельный диапазон: при numId вида 901 номера не отрисовываются вовсе —
# просмотрщик, судя по поведению, держит определения массивом по индексу
# (замер 03.08.2026: id 901 — пусто, id 10 — работает).
_NUM_BASE_ID = 0


def _numbering_root(doc):
    return doc.part.numbering_part.element


def _find_decimal_abstract(root):
    """Шаблонное определение десятичного списка «1.» — донор для клонирования."""
    for abstract in root.findall(qn("w:abstractNum")):
        if abstract.get(qn("w:abstractNumId")) in (None, ""):
            continue
        lvl = abstract.find(qn("w:lvl"))
        if lvl is None:
            continue
        fmt = lvl.find(qn("w:numFmt"))
        text = lvl.find(qn("w:lvlText"))
        style = lvl.find(qn("w:pStyle"))
        if (fmt is not None and fmt.get(qn("w:val")) == "decimal"
                and text is not None and text.get(qn("w:val")) == "%1."
                and style is None):          # без привязки к стилю
            return abstract
    return None


def _new_num_def(doc, *, left_tw=709, hanging_tw=340, start=1, font=None):
    """Независимое определение нумерации; вернуть numId.

    Определение НЕ собирается с нуля, а клонируется из шаблонного: собранное
    вручную рендерер молча игнорирует, потому что в нем нет w:nsid и w:tmpl
    (замер 03.08.2026 — номера не появлялись ни в одном просмотрщике, хотя
    XML был схемно верным). Клон наследует рабочую обвязку, меняются только
    отступы, шрифт и начальный номер.

    left_tw / hanging_tw — твипы (1 см = 567). Номер садится на left - hanging,
    текст на left.
    """
    from copy import deepcopy

    root = _numbering_root(doc)
    donor = _find_decimal_abstract(root)
    if donor is None:
        raise RuntimeError(
            "в numbering.xml нет десятичного определения-донора: шаблон "
            "python-docx изменился, автонумерация собрана не будет")

    used_abstract = [int(a.get(qn("w:abstractNumId")))
                     for a in root.findall(qn("w:abstractNum"))
                     if (a.get(qn("w:abstractNumId")) or "").isdigit()]
    used_num = [int(n.get(qn("w:numId"))) for n in root.findall(qn("w:num"))]
    new_id = max(used_abstract + used_num + [_NUM_BASE_ID]) + 1

    abstract = deepcopy(donor)
    abstract.set(qn("w:abstractNumId"), str(new_id))
    nsid = abstract.find(qn("w:nsid"))
    if nsid is not None:                      # уникальный, иначе Word склеит списки
        nsid.set(qn("w:val"), f"{0x7A000000 + new_id:08X}")
    tmpl = abstract.find(qn("w:tmpl"))
    if tmpl is not None:
        tmpl.set(qn("w:val"), f"{0x5B000000 + new_id:08X}")

    lvl = abstract.find(qn("w:lvl"))
    start_el = lvl.find(qn("w:start"))
    if start_el is not None:
        start_el.set(qn("w:val"), str(start))

    pPr = lvl.find(qn("w:pPr"))
    if pPr is None:
        pPr = OxmlElement("w:pPr")
        lvl.append(pPr)
    tabs = pPr.find(qn("w:tabs"))
    if tabs is not None:
        for tab in tabs.findall(qn("w:tab")):
            tab.set(qn("w:pos"), str(left_tw))
    ind = pPr.find(qn("w:ind"))
    if ind is None:
        ind = OxmlElement("w:ind")
        pPr.append(ind)
    ind.set(qn("w:left"), str(left_tw))
    ind.set(qn("w:hanging"), str(hanging_tw))

    rPr = lvl.find(qn("w:rPr"))
    if rPr is None:
        rPr = OxmlElement("w:rPr")
        lvl.append(rPr)
    rf = rPr.find(qn("w:rFonts"))
    if rf is None:
        rf = OxmlElement("w:rFonts")
        rPr.insert(0, rf)
    for attr in ("w:ascii", "w:hAnsi", "w:cs"):
        rf.set(qn(attr), font or FONT_BODY)
    rf.set(qn("w:hint"), "default")

    root.findall(qn("w:abstractNum"))[-1].addnext(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(new_id))
    ref = OxmlElement("w:abstractNumId")
    ref.set(qn("w:val"), str(new_id))
    num.append(ref)
    root.append(num)
    return new_id


def _apply_num(paragraph, num_id):
    pPr = paragraph._p.get_or_add_pPr()
    numPr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    numId = OxmlElement("w:numId")
    numId.set(qn("w:val"), str(num_id))
    numPr.append(ilvl)
    numPr.append(numId)
    # numPr в w:pPr идет сразу после pStyle — порядок задан схемой
    style = pPr.find(qn("w:pStyle"))
    if style is not None:
        style.addnext(numPr)
    else:
        pPr.insert(0, numPr)
    return paragraph


_MANUAL_NUM = __import__("re").compile(r"^\s*\d{1,2}\s*[.)]\s+")


def _strip_manual_number(text):
    """Снять номер, набранный руками: теперь его ставит Word.

    Вызовы вида add_request_item("1. Взыскать...") остаются рабочими — иначе
    в документе получилось бы «1. 1. Взыскать».
    """
    return _MANUAL_NUM.sub("", text, count=1)

class DocBuilder:
    """Построитель судебного документа по эталону шаблон.docx."""

    def __init__(self):
        self.doc = Document()
        section = self.doc.sections[0]
        section.page_width    = Mm(210)
        section.page_height   = Mm(297)
        section.top_margin    = Mm(20)
        section.bottom_margin = Mm(30)   # зона штампа экспедиции суда (эталон DOCX_FORMATTING.md)
        section.left_margin   = Mm(30)
        section.right_margin  = Mm(15)
        _set_docdefaults_line_spacing(self.doc)
        # Удалить дефолтный пустой параграф
        for p in self.doc.paragraphs:
            p._element.getparent().remove(p._element)
        # Определения нумерации создаются при первом обращении: документ без
        # приложений не должен тащить лишние записи в numbering.xml.
        self._num = {}

    def _num_id(self, kind):
        if kind not in self._num:
            spec = {
                "request":  dict(left_tw=709, hanging_tw=340),
                "appendix": dict(left_tw=709, hanging_tw=340),
                "body":     dict(left_tw=624, hanging_tw=624),
            }[kind]
            self._num[kind] = _new_num_def(self.doc, **spec)
        return self._num[kind]

    def add_empty(self):
        """Пустой параграф-разделитель."""
        self.doc.add_paragraph()

    def add_page_break(self):
        """Разрыв страницы. Приложения-фотографии начинаются с новой страницы,
        иначе снимок разрывает абзац по месту и лист печатается наполовину."""
        p = self.doc.add_paragraph()
        p.add_run().add_break(WD_BREAK.PAGE)
        return p

    def add_image(self, path, caption=None, width_mm=155):
        """Изображение по центру страницы, под ним подпись 11pt.

        width_mm по умолчанию 155: рабочее поле листа при левом поле 35 мм
        (L3) составляет 160 мм, запас в 5 мм не даёт Word ужать картинку
        и сорвать выравнивание по центру.
        """
        p = self.doc.add_paragraph()
        # Абзацный отступ намеренно НЕ задается: центрированный снимок и его
        # подпись отступа первой строки не имеют, а document_guard пропускает
        # абзацы с неустановленным отступом.
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after  = Pt(2)
        p.add_run().add_picture(path, width=Mm(width_mm))
        if caption:
            c = self.doc.add_paragraph()
            c.alignment = WD_ALIGN_PARAGRAPH.CENTER
            c.paragraph_format.space_after = Pt(10)
            _set_font(c.add_run(caption), 11)
        return p

    def add_title(self, text):
        """Главный заголовок документа: 14pt, bold, CENTER, sb=6, sa=4."""
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after  = Pt(4)
        _set_font(p.add_run(text), 14, bold=True, family=FONT_TITLE)
        _track(p, TRACK_TITLE_PT)
        return p

    def add_subtitle(self, text):
        """Подзаголовок: 13pt, bold, CENTER, sa=4."""
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(4)
        _set_font(p.add_run(text), 13, bold=True, family=FONT_TITLE)
        _track(p, TRACK_SUBTITLE_PT)
        return p

    def add_header_date(self, text):
        """Дата/время шапки: 13pt, bold, CENTER, sa=12."""
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(12)
        _set_font(p.add_run(text), 13, bold=True, family=FONT_DISPLAY)
        return p

    def add_section(self, text):
        """Заголовок раздела I/II/III: 12pt, bold, JUSTIFY, sb=12, sa=6."""
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after  = Pt(6)
        _set_font(p.add_run(text), 12, bold=True, family=FONT_DISPLAY)
        return p

    def add_subsection(self, text):
        """Нумерованный подраздел: 12pt, bold, JUSTIFY, sb=6, sa=3."""
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after  = Pt(3)
        _set_font(p.add_run(text), 12, bold=True, family=FONT_DISPLAY)
        return p

    def add_body(self, parts):
        """
        Основной текст: 12pt, JUSTIFY, fi=1.25cm, sa=6.
        parts: str или list[(text, bold)].
        """
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_before       = Pt(0)
        p.paragraph_format.space_after        = Pt(6)
        p.paragraph_format.first_line_indent  = Cm(1.25)
        if isinstance(parts, str):
            _set_font(p.add_run(parts), 12)
        else:
            for text, bold in parts:
                _set_font(p.add_run(text), 12, bold=bold)
        return p

    def add_body_spaced(self, parts):
        """Текст после маркированного списка: sb=6 sa=6 fi=1.25cm."""
        p = self.add_body(parts)
        p.paragraph_format.space_before = Pt(6)
        return p

    def add_numbered_body(self, parts):
        """Абзац тела со СКВОЗНЫМ номером, который ведет Word.

        Дает точную ссылку «пункт 14» вместо «страница 6»: страница едет от
        любой правки, номер пункта — нет. Нумерация сквозная по всему
        документу и продолжается через заголовки разделов и таблицы.
        Удаление, вставка и перестановка пункта пересчитываются программой.

        parts: str либо list[(text, bold)] — как в add_body.
        Отступы задает определение нумерации, вручную их не трогать.
        """
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_after = Pt(6)
        _apply_num(p, self._num_id("body"))
        if isinstance(parts, str):
            parts = [(parts, False)]
        for item in parts:
            text, bold = (item if isinstance(item, (tuple, list)) else (item, False))
            _set_font(p.add_run(text), 12, bold=bold)
        return p

    def add_bullet(self, text):
        """Маркированный список: List Bullet, JUSTIFY (как тело), li=1.5cm, sb=2, sa=2."""
        p = self.doc.add_paragraph(style="List Bullet")
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(2)
        p.paragraph_format.left_indent  = Cm(1.50)
        _set_font(p.add_run(text), 12)
        return p

    def add_table(self, headers, rows, widths=None, font_size=12,
                  aligns=None, borders=True):
        """
        Расчетная таблица (баланс, конвертация, сопоставление позиций сторон).

        Единственный штатный способ подать расчет таблицей. Маркированный список
        вместо таблицы — потеря структуры расчета (урок 02.08.2026, дело
        боевое дело, замечание К2 doc-reviewer).

        headers: list[str] — шапка, всегда bold.
        rows:    list[list] — ячейка либо str, либо (text, bold) для итоговых строк.
        widths:  list[int] | None — доли ширины колонок (нормируются на ширину
                 полосы набора: страница минус поля); None — колонки равные.
        font_size: 12 по эталону; 11 допустим для широких числовых таблиц
                 (набор размеров документа 14/13/12/11 не расширяется).
        aligns:  list из 'l'|'c'|'r' по колонкам; None — первая 'l', прочие 'r'.
        borders: True — тонкие границы 0.5 pt по всем ячейкам; False — без границ.
        """
        n_cols = len(headers)
        if aligns is None:
            aligns = ["l"] + ["r"] * (n_cols - 1)
        amap = {
            "l": WD_ALIGN_PARAGRAPH.LEFT,
            "c": WD_ALIGN_PARAGRAPH.CENTER,
            "r": WD_ALIGN_PARAGRAPH.RIGHT,
        }

        table = self.doc.add_table(rows=1 + len(rows), cols=n_cols)
        tbl = table._tbl
        tblPr = tbl.find(qn("w:tblPr"))
        if tblPr is None:
            tblPr = OxmlElement("w:tblPr")
            tbl.insert(0, tblPr)

        tblBorders = OxmlElement("w:tblBorders")
        for side in ["top", "left", "bottom", "right", "insideH", "insideV"]:
            el = OxmlElement(f"w:{side}")
            if borders:
                el.set(qn("w:val"), "single")
                el.set(qn("w:sz"), "4")          # 0.5 pt
                el.set(qn("w:space"), "0")
                el.set(qn("w:color"), "auto")
            else:
                el.set(qn("w:val"), "none")
            tblBorders.append(el)
        tblPr.append(tblBorders)

        # Ширина полосы набора в twips (учитывает поля секции, в т.ч. 35 мм для L3)
        section = self.doc.sections[0]
        avail = int(section.page_width - section.left_margin - section.right_margin) // 635
        if widths is None:
            widths = [1] * n_cols
        total = sum(widths)
        cols_tw = [max(600, int(avail * w / total)) for w in widths]

        tblW = OxmlElement("w:tblW")
        tblW.set(qn("w:w"), str(sum(cols_tw)))
        tblW.set(qn("w:type"), "dxa")
        tblPr.append(tblW)
        tblLayout = OxmlElement("w:tblLayout")
        tblLayout.set(qn("w:type"), "fixed")
        tblPr.append(tblLayout)

        tblGrid = OxmlElement("w:tblGrid")
        for w in cols_tw:
            col = OxmlElement("w:gridCol")
            col.set(qn("w:w"), str(w))
            tblGrid.append(col)
        tbl.insert(1, tblGrid)

        def _fill(cell, value, col_idx, bold_default=False):
            cell.paragraphs[0].clear()
            p = cell.paragraphs[0]
            p.alignment = amap[aligns[col_idx]]
            p.paragraph_format.space_before      = Pt(2)
            p.paragraph_format.space_after       = Pt(2)
            p.paragraph_format.first_line_indent = Cm(0)
            text, bold = (value if isinstance(value, tuple) else (value, bold_default))
            # Гарнитура одна на всё (решение владельца 04.08.2026). Роли
            # различаются начертанием: шапка таблицы — полужирный, числовые
            # колонки идут вправо и держат разряды выравниванием, а не моноширинностью.
            if bold_default:
                family = FONT_DISPLAY
            elif aligns[col_idx] == "r":
                family = FONT_MONO
            else:
                family = FONT_BODY
            _set_font(p.add_run(text), font_size, bold=bold, family=family)

        for j, head in enumerate(headers):
            _fill(table.rows[0].cells[j], head, j, bold_default=True)
        for i, row in enumerate(rows):
            for j, value in enumerate(row):
                _fill(table.rows[1 + i].cells[j], value, j)

        return table

    def add_quote(self, text):
        """Блок-цитата (дословная норма/формула): JUSTIFY, отступы 1.25 см
        слева и справа, без первой строки, 11pt, sb=6, sa=6.

        Эталон CONTENT_DESIGN.md: единственный правильный контейнер для
        дословного текста нормы закона, судебного акта или формулы (40+ слов).
        Отступ служит маркером цитаты — кавычки не нужны.
        """
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.left_indent       = Cm(1.25)
        p.paragraph_format.right_indent      = Cm(1.25)
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.space_before      = Pt(6)
        p.paragraph_format.space_after       = Pt(6)
        _set_font(p.add_run(text), 11)
        if not hasattr(self, "_quote_paragraphs"):
            self._quote_paragraphs = set()
        self._quote_paragraphs.add(p._p)
        return p

    def add_proshyu(self):
        """Заголовок ПРОШУ: CENTER, bold, 12pt, sa=6."""
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(6)
        _set_font(p.add_run("ПРОШУ:"), 12, bold=True, family=FONT_DISPLAY)
        return p

    def add_request_item(self, text):
        """Пункт просительной части: автонумерация Word, JUSTIFY, sb=2, sa=2.

        Номер ставит Word (w:numPr), а не составитель: удаление или вставка
        пункта пересчитывает остальные сама. Номер, набранный в тексте
        («1. Взыскать...»), снимается — иначе задвоится.
        """
        p = self.doc.add_paragraph(style="List Paragraph")
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(2)
        _apply_num(p, self._num_id("request"))
        _set_font(p.add_run(_strip_manual_number(text)), 12)
        return p

    def add_appendices(self):
        """Заголовок ПРИЛОЖЕНИЯ: JUSTIFY, bold, 12pt, sb=12, sa=6."""
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after  = Pt(6)
        _set_font(p.add_run("ПРИЛОЖЕНИЯ:"), 12, bold=True, family=FONT_DISPLAY)
        return p

    def add_appendix_item(self, text):
        """Пункт приложений: автонумерация Word, JUSTIFY, sa=12.

        Приложения правятся чаще всего — здесь ручной пересчет болел сильнее
        всего. Номер ведет Word; ссылки в тексте («Приложение 3») остаются
        на совести составителя, их сверяет document_guard.py.
        """
        p = self.doc.add_paragraph(style="List Paragraph")
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(12)
        _apply_num(p, self._num_id("appendix"))
        _set_font(p.add_run(_strip_manual_number(text)), 12)
        return p

    def add_signature(self, name, date, gap_spaces=40):
        """
        Строка подписи: ФИО + пробелы + дата в одном параграфе.
        RIGHT, sb=6, sa=6. (Правило форматирования: подпись справа.)
        """
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after  = Pt(6)
        _set_font(p.add_run(name), 12)
        _set_font(p.add_run(" " * gap_spaces), 12)
        _set_font(p.add_run(date), 12)
        return p

    def add_signature_table(self, role, name, date=None):
        """
        Блок подписи 3-колоночной таблицей без границ (ГОСТ Р 7.0.97-2016 п. 5.22).

        [Роль/Должность]   [пробел для подписи]   [И. О. Фамилия]
        Устойчиво к открытию в LibreOffice и Google Docs (в отличие от подписи
        пробелами). Дата (если задана) — отдельным параграфом под таблицей, LEFT.
        """
        table = self.doc.add_table(rows=1, cols=3)

        tbl = table._tbl
        tblPr = tbl.find(qn("w:tblPr"))
        if tblPr is None:
            tblPr = OxmlElement("w:tblPr")
            tbl.insert(0, tblPr)
        tblBorders = OxmlElement("w:tblBorders")
        for side in ["top", "left", "bottom", "right", "insideH", "insideV"]:
            el = OxmlElement(f"w:{side}")
            el.set(qn("w:val"), "none")
            tblBorders.append(el)
        tblPr.append(tblBorders)

        tblGrid = OxmlElement("w:tblGrid")
        for w in ["5000", "2339", "2000"]:
            col = OxmlElement("w:gridCol"); col.set(qn("w:w"), w)
            tblGrid.append(col)
        tbl.insert(1, tblGrid)

        row = table.rows[0]
        c0 = row.cells[0]
        c0.paragraphs[0].clear()
        p0 = c0.paragraphs[0]
        p0.alignment = WD_ALIGN_PARAGRAPH.LEFT
        _set_font(p0.add_run(role), 12)
        c1 = row.cells[1]
        c1.paragraphs[0].clear()
        c1.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        c2 = row.cells[2]
        c2.paragraphs[0].clear()
        p2 = c2.paragraphs[0]
        p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        _set_font(p2.add_run(name), 12)

        if date:
            pd = self.doc.add_paragraph()
            pd.alignment = WD_ALIGN_PARAGRAPH.LEFT
            pd.paragraph_format.space_before = Pt(6)
            pd.paragraph_format.space_after  = Pt(6)
            _set_font(pd.add_run(date), 12)
        return table

    def add_final_empty(self):
        """Последний пустой параграф: JUSTIFY, sb=6, sa=0."""
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after  = Pt(0)
        return p

    def add_addressee_table(self, blocks):
        """
        Адресная шапка досудебного документа (КОМУ / ОТ) — плавающая таблица
        в правой части листа, без границ. В отличие от add_header_table не
        создаёт служебных строк суда и дела.

        blocks: list[dict] с ключами:
            label: «КОМУ:» / «ОТ:»
            lines: list[(text, bold)]
        """
        table = self.doc.add_table(rows=len(blocks), cols=2)

        tbl = table._tbl
        tblPr = tbl.find(qn("w:tblPr"))
        if tblPr is None:
            tblPr = OxmlElement("w:tblPr")
            tbl.insert(0, tblPr)

        tblBorders = OxmlElement("w:tblBorders")
        for side in ["top", "left", "bottom", "right", "insideH", "insideV"]:
            el = OxmlElement(f"w:{side}")
            el.set(qn("w:val"), "none")
            tblBorders.append(el)
        tblPr.append(tblBorders)

        tblGrid = OxmlElement("w:tblGrid")
        col1 = OxmlElement("w:gridCol"); col1.set(qn("w:w"), "1400")
        col2 = OxmlElement("w:gridCol"); col2.set(qn("w:w"), "7939")
        tblGrid.append(col1); tblGrid.append(col2)
        tbl.insert(1, tblGrid)

        tblpPr = OxmlElement("w:tblpPr")
        tblpPr.set(qn("w:horzAnchor"), "margin")
        tblpPr.set(qn("w:vertAnchor"), "text")
        tblpPr.set(qn("w:tblpY"), "-325")
        tblPr.insert(0, tblpPr)

        for i, block in enumerate(blocks):
            row = table.rows[i]
            lc = row.cells[0]
            lc.paragraphs[0].clear()
            lp = lc.paragraphs[0]
            lp.alignment = WD_ALIGN_PARAGRAPH.LEFT
            _set_font(lp.add_run(block["label"]), 12, bold=True, family=FONT_DISPLAY)
            rc = row.cells[1]
            rc.paragraphs[0].clear()
            first = True
            for text, bold in block["lines"]:
                if first:
                    p = rc.paragraphs[0]
                    first = False
                else:
                    p = rc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                _set_font(p.add_run(text), 12, bold=bold)
        return table

    def add_header_table(self, court_name, court_route, parties, case_number, instance=None):
        """
        Шапка-реквизиты: плавающая таблица 2×N без границ.

        court_name: str — «ВЕРХОВНЫЙ СУД РЕСПУБЛИКИ ТАТАРСТАН»
        court_route: str — «через Вахитовский районный суд г. Казани»
        parties: list[dict] с ключами:
            label: «ИСТЕЦ:» / «ОТВЕТЧИКИ:»
            lines: list[(text, bold)]
        case_number: str — «Дело № 2-5612/2025»
        instance: str — «Суд первой инстанции: ...» (опционально)
        """
        rows_count = 1 + len(parties) + 1  # суд + стороны + дело
        table = self.doc.add_table(rows=rows_count, cols=2)

        # Убрать все границы
        tbl = table._tbl
        tblPr = tbl.find(qn("w:tblPr"))
        if tblPr is None:
            tblPr = OxmlElement("w:tblPr")
            tbl.insert(0, tblPr)

        tblBorders = OxmlElement("w:tblBorders")
        for side in ["top", "left", "bottom", "right", "insideH", "insideV"]:
            el = OxmlElement(f"w:{side}")
            el.set(qn("w:val"), "none")
            tblBorders.append(el)
        tblPr.append(tblBorders)

        # Ширины колонок: доли полосы набора (38/62), чтобы шапка не вылезала
        # за поле при левом поле 35 мм у документов L3
        section = self.doc.sections[0]
        avail = int(section.page_width - section.left_margin - section.right_margin) // 635
        tblGrid = OxmlElement("w:tblGrid")
        col1 = OxmlElement("w:gridCol"); col1.set(qn("w:w"), str(int(avail * 0.38)))
        col2 = OxmlElement("w:gridCol"); col2.set(qn("w:w"), str(avail - int(avail * 0.38)))
        tblGrid.append(col1); tblGrid.append(col2)
        tbl.insert(1, tblGrid)

        # Плавающая позиция (floating)
        tblpPr = OxmlElement("w:tblpPr")
        tblpPr.set(qn("w:horzAnchor"), "margin")
        tblpPr.set(qn("w:vertAnchor"), "text")
        tblpPr.set(qn("w:tblpY"), "-325")
        tblPr.insert(0, tblpPr)

        # Строка 0: суд
        row0 = table.rows[0]
        row0.cells[0].text = ""
        c = row0.cells[1]
        c.paragraphs[0].clear()
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        _set_font(p.add_run(court_name), 12, bold=True, family=FONT_TITLE)
        if court_route:
            p2 = c.add_paragraph()
            p2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            _set_font(p2.add_run(court_route), 12)

        # Строки сторон
        for i, party in enumerate(parties):
            row = table.rows[1 + i]
            # Метка слева
            lc = row.cells[0]
            lc.paragraphs[0].clear()
            lp = lc.paragraphs[0]
            lp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            _set_font(lp.add_run(party["label"]), 12, bold=True, family=FONT_DISPLAY)
            # Реквизиты справа
            rc = row.cells[1]
            rc.paragraphs[0].clear()
            first = True
            for text, bold in party["lines"]:
                if first:
                    p = rc.paragraphs[0]
                    first = False
                else:
                    p = rc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                _set_font(p.add_run(text), 12, bold=bold)

        # Последняя строка: дело
        last_row = table.rows[-1]
        last_row.cells[0].text = ""
        c = last_row.cells[1]
        c.paragraphs[0].clear()
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        _set_font(p.add_run(case_number), 12)
        if instance:
            p2 = c.add_paragraph()
            p2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            _set_font(p2.add_run(instance), 12)

        return table

    def add_page_numbers(self, hide_on_first=True):
        """
        Номера страниц: арабские цифры, центр НИЖНЕГО поля, первая страница без
        номера. Решение владельца 04.08.2026 — номер переехал из верхнего поля
        вниз. Обязательно для документов 2+ страниц.

        Штамп суда о приеме ставят в правом нижнем углу — центр нижнего поля
        его не занимает.

        Вручную звать не обязательно: `save()` вызывает сам, если не вызвали.
        """
        from copy import deepcopy

        self._paginated = True

        section = self.doc.sections[0]
        section.different_first_page_header_footer = hide_on_first
        footer = section.footer
        footer.is_linked_to_previous = False
        # Номеру нужен СВОЙ абзац. Нижний колонтитул может быть уже занят —
        # в образцах legal design туда пишут сноску («Образец, фактура вымышлена»),
        # и номер, подсаженный в тот же абзац, вставал впритык к чужому тексту
        # по центру. Занятый абзац не трогаем, берем следующий; абзац, где номер
        # уже стоит, переиспользуем — иначе повторный вызов даст два номера.
        for cand in footer.paragraphs:
            if "PAGE" in cand._p.xml:
                return cand          # номер уже стоит — второй не ставим
        p = next((c for c in footer.paragraphs
                  if not c.text.strip() and not c.runs), None)
        if p is None:
            p = footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        probe = p.add_run("")
        _set_font(probe, 12, family=FONT_DISPLAY)
        rPr = probe._r.find(qn("w:rPr"))

        fld = OxmlElement("w:fldSimple")
        fld.set(qn("w:instr"), " PAGE ")
        r = OxmlElement("w:r")
        if rPr is not None:
            r.append(deepcopy(rPr))
        text_el = OxmlElement("w:t")
        text_el.text = "2"
        r.append(text_el)
        fld.append(r)
        p._p.append(fld)
        return p

    def _strip_yo(self):
        """Заменяет ё→е, Ё→Е во всех текстовых узлах (правило проекта «нет ё»)."""
        for t in self.doc.element.iter(qn("w:t")):
            if t.text and ("ё" in t.text or "Ё" in t.text):
                t.text = t.text.replace("ё", "е").replace("Ё", "Е")

    # Категории scan_legal.sh, при которых документ не выпускается. Мягкие
    # (хеджирование, номинализации, ритм) только предупреждают — их снимает
    # сам скилл, а формальным блокером они дают ложные срабатывания.
    HUMANIZER_BLOCKERS = (
        "HARD BANS",
        "L18 плейсхолдер",
        "L18 артефакт копипасты",
        "L18 невидимые символы",
        "L18 латиница в кириллице",
    )

    def _document_text(self):
        """Весь видимый текст документа: параграфы плюс ячейки таблиц.

        Блок-цитаты отдаются с markdown-префиксом «> ». Без него гейт читает
        дословную норму закона как авторский текст и бракует документ за чужие
        слова: scan_legal.sh выводит цитаты из авторских категорий именно по
        этому префиксу, а в .docx разметки нет — цитату задает отступ.
        Прецедент 04.08.2026: ч. 3 ст. 11 УПК РФ («достаточных данных») не
        сохранялась, хотя тот же текст в .md проходил чисто.
        """
        quotes = getattr(self, "_quote_paragraphs", set())
        parts = ["> " + p.text if p._p in quotes else p.text
                 for p in self.doc.paragraphs]
        for t in self.doc.tables:
            for row in t.rows:
                parts.extend(c.text for c in row.cells)
        # Абзацы разделяются пустой строкой: scan_legal.sh склеивает соседние
        # непустые строки в один абзац по правилам markdown, и цитата теряет
        # свой префикс «> », если предыдущий абзац стоит вплотную.
        return "\n\n".join(x for x in parts if x.strip())

    def _matches_approved(self, md_path):
        """Собранный текст обязан быть равен одобренной редакции .md — целиком.

        Сравнение после нормализации: разметка markdown, регистр и «ё» (её
        сборщик стрипает сам) на равенство не влияют. Одобренный текст обязан
        войти в документ целиком и подряд — и документ обязан ИСЧЕРПЫВАТЬСЯ
        одобренным текстом: «вхождение подстроки» позволяло ДОПИСАТЬ в конец
        новое требование («обратить взыскание на квартиру ответчика»), и
        сборка проходила — Кони видел меньше, чем уходило в суд (этап 9.19,
        круг 7). После одобренного текста допустим только хвост-подпись
        (роль/ФИО и дата сборки), до него — ничего. Вернет True/False.
        """
        import re as _re
        from pathlib import Path as _Path

        def norm(s):
            s = _re.sub(r"[*_`#>|\-]+", " ", s)
            s = s.replace("ё", "е").replace("Ё", "е")
            return _re.sub(r"\s+", " ", s).strip().lower()

        def tolko_podpis(hvost):
            """Хвост после одобренного текста — только подпись: несколько слов
            (роль/ФИО) и дата ДД.ММ.ГГГГ. Дописанное требование подписью не
            выглядит: оно длиннее и содержит сказуемое, а не только реквизит."""
            hvost = hvost.strip()
            if not hvost:
                return True
            if len(hvost.split()) > 6:
                return False
            return bool(_re.match(
                r"^[a-zа-яё\d\s.\-]*\d{1,2}\.\d{1,2}\.\d{4}$", hvost))

        try:
            md_norm = norm(_Path(md_path).read_text(encoding="utf-8", errors="replace"))
        except OSError:
            return False
        if not md_norm:
            return True
        doc_norm = norm(self._document_text())
        idx = doc_norm.find(md_norm)
        if idx < 0:
            return False
        if doc_norm[:idx].strip():
            return False          # перед одобренным текстом что-то дописано
        return tolko_podpis(doc_norm[idx + len(md_norm):])

    def _humanizer_gate(self):
        """Прогнать текст через scan_legal.sh. Вернуть список сработавших
        блокирующих категорий; пустой список — документ чист.

        Скрипт недоступен — гейт не срабатывает, но об этом печатается
        предупреждение: молчаливое отключение проверки хуже ее отсутствия.
        """
        import subprocess
        from pathlib import Path as _P

        scan = _P.home() / ".claude/skills/humanizer-legal/scripts/scan_legal.sh"
        if not scan.exists():
            print(f"ВНИМАНИЕ: {scan} не найден — проверка humanizer-legal не выполнена.")
            return []
        try:
            out = subprocess.run(
                ["bash", str(scan), "-"], input=self._document_text(),
                capture_output=True, text=True, timeout=60).stdout
        except (OSError, subprocess.SubprocessError) as e:
            print(f"ВНИМАНИЕ: humanizer-legal не отработал ({e}) — проверка не выполнена.")
            return []

        hits = []
        for line in out.splitlines():
            parts = line.strip().split(None, 1)
            if len(parts) == 2 and parts[0].isdigit() and int(parts[0]) > 0:
                if parts[1] in self.HUMANIZER_BLOCKERS:
                    hits.append(f"{parts[1]} ({parts[0]})")
        return hits

    def save(self, path):
        """Сохранить документ. Перед записью — авто-стрип буквы ё и гейт
        humanizer-legal: документ со следами автогенерации или незаполненным
        плейсхолдером в суд не выпускается.

        Плюс неизменяемый снимок-черновик в `_baselines/` рядом — база «ДО» для
        самообучения по правкам доверителя (redline). Снимок = последняя выданная
        версия; правки доверителя сравниваются именно с ней.

        Guard (урок 14.07.2026): если выданный файл отличается от baseline —
        его правил доверитель; перезапись вслепую запрещена безусловно
        (этап 9: env-обход снят, сторож не должен уметь себя выключать
        переменной, о которой claude_guard не знает).
        """
        import filecmp
        import os
        import shutil
        from pathlib import Path as _P

        self._strip_yo()
        p = _P(path)
        # `.agent/drafts` — ДВЕ составляющие пути, и проверка `in p.parts` на ней
        # всегда ложна. Раскладку спрашиваем у контракта, а не у литерала.
        sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
        import case_paths as _cp
        parts = p.parts
        # снимок только для реальных документов дел (не тест/tmp), без рекурсии
        is_case_doc = (p.parent.name != "_baselines"
                       and ("cases" in parts or _cp.READY in parts
                            or (_cp.AGENT_DIR in parts and "drafts" in parts)))

        # Точка сборки одна: готовый документ живёт в GOTOVO/, снимок — рядом с
        # черновиком, в .agent/drafts/_baselines/. Иначе снимки засоряют папку,
        # за которой человек приходит за готовым.
        case_root = None
        for i, part in enumerate(parts):
            if part == _cp.READY or (part == _cp.AGENT_DIR and "drafts" in parts[i:]):
                case_root = _P(*parts[:i])
                break
        bfile = (_cp.baselines(case_root) / p.name) if case_root \
            else (p.parent / "_baselines" / p.name)

        # Вердикт Кони привязан к SHA-256 редакции .md. Решение владельца:
        # «.docx собирается один раз, после вердикта Кони» — поэтому гейт
        # стоит на ЛЮБОМ документе дела, а не только на сегменте GOTOVO:
        # прежде .docx в .agent/drafts/ собирался вообще без вердикта и без
        # парного .md (проба круга 6, 20.08.2026).
        if is_case_doc:
            import verdict as _v
            md = None
            kandidaty = []
            if case_root is not None and _cp.READY in parts:
                kandidaty.append(_cp.drafts(case_root) / (p.stem + ".md"))
            kandidaty.append(p.with_suffix(".md"))
            for cand in kandidaty:
                if cand.is_file():
                    md = cand
                    break
            if md is None:
                print(f"СТОП, НЕ СОХРАНЕНО: {p.name} — не найден парный черновик .md "
                      f"в {_cp.DRAFTS}/. Документ собирается ИЗ черновика, "
                      f"одобренного Кони, а не из воздуха.")
                return
            problems = _v.check(md)
            if problems:
                print("СТОП, НЕ СОХРАНЕНО: сборка .docx запрещена вердиктом.")
                for x in problems:
                    print("  · " + x)
                print("  Провести раунд Кони и записать вердикт: "
                      "python3 scripts/verdict.py ЧЕРНОВИК.md --record --verdict "
                      "'ГОТОВ К ПОДАЧЕ' -r N")
                return
            # Отпечаток .md доказывает лишь неизменность .md — а не то, что в
            # .docx собрали ИМЕННО одобренный текст. Проба круга 6: Кони
            # одобрил «взыскать 100 000 (сто тысяч) рублей задолженности», а
            # в суд собрали «5 000 000 и обратить взыскание на квартиру» —
            # сборка прошла. Собранный документ обязан содержать одобренную
            # редакцию дословно (после нормализации разметки и «ё») — и
            # исчерпываться ею: дописать требование в конец нельзя, допустим
            # только хвост-подпись (круг 7, этап 9.19).
            if not self._matches_approved(md):
                print(f"СТОП, НЕ СОХРАНЕНО: собранный текст не совпадает с "
                      f"одобренной редакцией {md.name}. Отпечаток вердикта "
                      f"привязан к .md — собирать нужно его, а не другой текст.")
                return

        if (is_case_doc and p.exists() and bfile.exists()
                and not filecmp.cmp(p, bfile, shallow=False)):
            print(f"СТОП, НЕ СОХРАНЕНО: {p} отличается от _baselines/ — "
                  f"вероятны правки доверителя, внесенные напрямую в файл. "
                  f"Сначала redline-разбор («изучи мои правки»).")
            return

        # Нумерация обязательна по ГОСТ Р 7.0.97 для документов 2+ страниц
        # (DOCX_FORMATTING.md:28). Требовать ее памятью составителя ненадежно:
        # на 02.08.2026 живого вызова не осталось ни на одном маршруте.
        if not getattr(self, "_paginated", False):
            self.add_page_numbers()

        # Гейт humanizer-legal вынесен из сборки в прогон по `.md` каждый раунд
        # (`scripts/verdict.py --scan`). На собранном `.docx` он срабатывал один
        # раз и слишком поздно — текст уже стал документом. Для документов вне
        # конвейера (тесты, разовые сборки) остаётся здесь как последний рубеж.
        if not is_case_doc:
            blockers = self._humanizer_gate()
            if blockers:
                print(f"СТОП, НЕ СОХРАНЕНО: документ не прошел humanizer-legal.\n"
                      f"  Сработали блокирующие категории: {', '.join(blockers)}.\n"
                      f"  Прогнать `python3 scripts/verdict.py ФАЙЛ.md --scan`, затем повторить.")
                return

        self.doc.save(path)
        print(f"Сохранено: {path}")
        if is_case_doc:
            try:
                bfile.parent.mkdir(parents=True, exist_ok=True)
                shutil.copy2(path, bfile)  # перезапись: baseline = свежая версия
            except OSError as e:
                # Fail-open закрыт: без снимка разбор правок доверителя сравнивает
                # документ сам с собой и молча ничему не учится. Прежде здесь стояло
                # предупреждение, которое никто не читал.
                raise RuntimeError(
                    f"снимок в {bfile} не записан ({e}). Документ сохранён, но база «ДО» "
                    f"для redline-разбора не обновлена — закройте файл в других "
                    f"приложениях и повторите save()") from e


def _verdict_gate_checks(tmp):
    """Жизненный цикл документа: .docx собирается ОДИН раз, из редакции, одобренной Кони."""
    import sys as _s
    from pathlib import Path as _P
    _s.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
    import case_paths as _cp
    import verdict as _v

    case = _P(tmp) / "cases" / "ivanov-ivan" / "delo-2026"
    _cp.drafts(case).mkdir(parents=True)
    _cp.ready(case).mkdir(parents=True)
    md = _cp.drafts(case) / "isk_v1.md"
    # Сборка обязана воспроизводить одобренный текст (этап 9, круг 6): тело
    # фикстуры равно тексту .md, иначе гейт равенства правомерно не пустит.
    md.write_text("# ИСКОВОЕ ЗАЯВЛЕНИЕ\n\nТекст документа без плейсхолдеров.\n",
                  encoding="utf-8")
    target = _cp.ready(case) / "isk_v1.docx"

    def build():
        b = DocBuilder()
        b.add_title("ИСКОВОЕ ЗАЯВЛЕНИЕ")
        b.add_body("Текст документа без плейсхолдеров.")
        return b

    build().save(str(target))
    no_verdict = not target.exists()

    _v.record(md, "ТРЕБУЕТ ПРАВОК", 1)
    build().save(str(target))
    not_ready = not target.exists()

    _v.record(md, _v.READY, 2)
    build().save(str(target))
    saved = target.exists()

    bl = _cp.baselines(case) / "isk_v1.docx"
    snapshot_right_place = bl.is_file() and not (_cp.ready(case) / "_baselines").exists()

    # Текст правится ПОСЛЕ одобрения — прежний вердикт не должен пускать сборку
    target.unlink()
    md.write_text("# Иск\n\nТекст.\n\nДописано после одобрения.\n", encoding="utf-8")
    build().save(str(target))
    stale_blocked = not target.exists()

    return [
        ("сборка без вердикта Кони запрещена", no_verdict),
        ("сборка при вердикте ТРЕБУЕТ ПРАВОК запрещена", not_ready),
        ("сборка по одобренной редакции проходит", saved),
        ("снимок лёг в .agent/drafts/_baselines, а не в GOTOVO", snapshot_right_place),
        ("правка .md после одобрения снова запрещает сборку", stale_blocked),
    ]


def selftest() -> int:
    """Проверка сборщика без сети: гарнитура и место номера страницы."""
    import tempfile
    from docx import Document

    tmp = tempfile.mkdtemp()

    plain = os.path.join(tmp, "plain.docx")
    b = DocBuilder()
    b.add_title("ИСКОВОЕ ЗАЯВЛЕНИЕ")
    b.add_section("I. ОБСТОЯТЕЛЬСТВА")
    b.add_body("Текст документа.")
    b.add_table(["Основание", "Сумма, руб."], [["Долг", "150 000,00"]])
    b.add_signature("Представитель", "04.08.2026")
    b.save(plain)

    # Колонтитул уже занят чужой строкой (образцы legal design пишут туда сноску).
    busy = os.path.join(tmp, "busy.docx")
    b2 = DocBuilder()
    b2.add_body("Текст документа.")
    f = b2.doc.sections[0].footer
    f.is_linked_to_previous = False
    p0 = f.paragraphs[0] if f.paragraphs else f.add_paragraph()
    p0.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _set_font(p0.add_run("Образец. Фактура вымышлена."), 11)
    b2.save(busy)

    twice = os.path.join(tmp, "twice.docx")
    b3 = DocBuilder()
    b3.add_body("Текст документа.")
    b3.add_page_numbers()
    b3.add_page_numbers()
    b3.save(twice)

    def fonts(path):
        d, out = Document(path), set()
        def walk(c):
            for par in getattr(c, "paragraphs", []):
                for r in par.runs:
                    if r.font.name:
                        out.add(r.font.name)
            for t in getattr(c, "tables", []):
                for row in t.rows:
                    for cell in row.cells:
                        walk(cell)
        walk(d)
        walk(d.sections[0].footer)
        return out

    d_plain, d_busy, d_twice = (Document(p) for p in (plain, busy, twice))
    foot_plain = d_plain.sections[0].footer
    foot_busy = d_busy.sections[0].footer

    checks = [
        ("гарнитура в документе одна — PT Serif", fonts(plain) == {FONT_BODY}),
        ("роли не разъезжаются по гарнитурам",
         FONT_TITLE == FONT_DISPLAY == FONT_MONO == FONT_BODY),
        ("номер страницы в нижнем колонтитуле", "PAGE" in foot_plain._element.xml),
        ("в верхнем колонтитуле номера нет",
         "PAGE" not in d_plain.sections[0].header._element.xml),
        ("номер по центру",
         any(p.alignment == WD_ALIGN_PARAGRAPH.CENTER and "PAGE" in p._p.xml
             for p in foot_plain.paragraphs)),
        ("первая страница без номера",
         d_plain.sections[0].different_first_page_header_footer),
        # Занятый колонтитул: номер обязан взять свой абзац, чужую строку не трогая.
        ("чужая строка колонтитула не перекроена",
         any(p.text.startswith("Образец") and p.alignment == WD_ALIGN_PARAGRAPH.LEFT
             and "PAGE" not in p._p.xml for p in foot_busy.paragraphs)),
        ("номер получил собственный абзац",
         any(not p.text.strip() and "PAGE" in p._p.xml for p in foot_busy.paragraphs)),
        ("повторный вызов не ставит второй номер",
         d_twice.sections[0].footer._element.xml.count(" PAGE ") == 1),
        ("буква ё вычищена при сохранении", "ё" not in "".join(
            p.text for p in d_plain.paragraphs)),
    ] + _verdict_gate_checks(tmp)

    for name, ok in checks:
        print(f"  {'✓' if ok else '✗'} {name}")
    bad = [n for n, ok in checks if not ok]
    if bad:
        print(f"selftest ПРОВАЛЕН: {len(bad)} из {len(checks)}")
        return 1
    print(f"selftest пройден: {len(checks)}/{len(checks)}")
    return 0


if __name__ == "__main__":
    if "--selftest" in sys.argv:
        sys.exit(selftest())
    # Тест-пример
    b = DocBuilder()
    b.add_header_table(
        court_name="ВЕРХОВНЫЙ СУД РЕСПУБЛИКИ ТАТАРСТАН",
        court_route="через Вахитовский районный суд г. Казани",
        parties=[
            {
                "label": "ИСТЕЦ:",
                "lines": [("Иванов Иван Иванович", True), ("адрес: г. Казань...", False)]
            },
            {
                "label": "ОТВЕТЧИК:",
                "lines": [("ООО «Ответчик»", True), ("ИНН: 1234567890", False)]
            },
        ],
        case_number="Дело № 2-0001/2026",
        instance="Суд первой инстанции: Вахитовский р.с. г. Казани"
    )
    b.add_empty()
    b.add_title("ИСКОВОЕ ЗАЯВЛЕНИЕ")
    b.add_subtitle("о взыскании неосновательного обогащения")
    b.add_header_date("27 мая 2026 года")
    b.add_empty()
    b.add_section("I. ОБСТОЯТЕЛЬСТВА ДЕЛА")
    b.add_body([("В 2024 году истец ", False), ("передал ответчику денежные средства", True), (" в сумме...", False)])
    b.add_section("II. ПРАВОВОЕ ОБОСНОВАНИЕ")
    b.add_subsection("1. Нормы применимого права")
    b.add_body("В соответствии с положениями гражданского законодательства...")
    b.add_proshyu()
    b.add_request_item("1. Взыскать с ответчика сумму неосновательного обогащения...")
    b.add_appendices()
    b.add_appendix_item("1. Квитанция об уплате государственной пошлины.")
    b.add_signature("Иванов Иван Иванович", "27.05.2026")
    b.add_final_empty()
    b.save("/tmp/test_doc.docx")
