#!/usr/bin/env python3
"""document_guard.py — формат .docx и согласованность .md/.docx машиной, не глазами.

Блок Б6 у проверяющего был инструкцией «распечатать поля и кегли и посмотреть» —
то есть сравнение со спецификацией делала модель, тратя токены на арифметику и
регулярно давая ложные тревоги. Здесь то же сравнение делает код.

Проверяет по .claude/skills/doc-drafter/DOCX_FORMATTING.md:
  поля 20/30/30/15 мм (L3 и кассация ВС — левое 35); ОДНА гарнитура PT Serif на
  весь документ (решение владельца 04.08.2026, иные гарнитуры недопустимы);
  кегли 14/13/12/11, межстрочный 1.15, абзацный отступ 1.25 см, тело JUSTIFY,
  отсутствие курсива и подчеркивания, номер страницы — БЕЗУСЛОВНО в каждом
  документе и БЕЗУСЛОВНО в нижнем колонтитуле.
Плюс: текст .docx совпадает с .md, приложения пронумерованы сквозно и каждое
упомянуто в тексте, запрещенная буква «ё», даты в формате ДД.ММ.ГГГГ.
Плюс (этап 9.8, требование владельца): денежная сумма пишется цифрами и тут
же прописью в КРУГЛЫХ скобках — «1 000 (одна тысяча) рублей» — и пропись
обязана СОВПАДАТЬ с числом. Разбор денег — единый на проект,
`scripts/money_rule.py` (словарь валют и числительных из `scripts/propis.py`,
дословные цитаты в елочках, сверка по шести падежам) — его же импортирует
verdict.py: две копии правила расходились, и целые классы документов
не выпускались (этап 9.22, круг 9). Прописи требует только сумма денег:
даты, статьи, номера дел, ИНН, проценты и листы дела ее не требуют.

    document_guard.py ДОКУМЕНТ.docx [--md ДОКУМЕНТ.md] [--l3]
    document_guard.py --selftest

Код возврата: 0 — чисто, 1 — есть нарушения.
"""
from __future__ import annotations

import argparse
import os
import re
import sys

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
try:
    import money_rule as _mr  # единый денежный разбор проекта (валюты,
except ImportError:           # числительные, цитаты, сверка прописи)
    _mr = None

SPEC_MARGINS_MM = {"top": 20, "bottom": 30, "left": 30, "right": 15}
SPEC_MARGINS_MM_L3 = {**SPEC_MARGINS_MM, "left": 35}
# Договор и прочие несудебные документы: нижнее поле 30 мм не нужно — зона
# штампа экспедиции суда там ни при чем, а поля симметричнее и читать удобнее.
# Шрифты, кегли, интервалы и нумерация страниц остаются общими для всего.
SPEC_MARGINS_MM_DOGOVOR = {"top": 20, "bottom": 20, "left": 25, "right": 20}
# Бланк договора адвокатского центра (DOCX_FORMATTING.md §8). Полоса набора 185 мм:
# уже — и таблицы фирменной шапки сжимаются, заголовки начинают рваться по слогам.
# Профиль установлен правкой владельца 20.08.2026 на договоре оказания услуг.
SPEC_MARGINS_MM_ADVOKAT = {"top": 20, "bottom": 20, "left": 15, "right": 10}
SPEC_SIZES_ADVOKAT = {7.0, 8.0, 9.0, 9.5, 10.0, 11.0, 12.0}
SPEC_LINE_SPACING_ADVOKAT = 1.0
# Бланк несет два отступа: 1,0 см у преамбул и 1,25 см у трех абзацев-продолжений
# внутри пунктов. Оба пришли из фирменного образца, править их — ломать бланк.
SPEC_INDENT_CM_ADVOKAT = {1.0, 1.25}
# ОДНА гарнитура на документ — PT Serif. Решение владельца 04.08.2026, отменяет
# набор из четырех гарнитур от 03.08.2026. PT Serif свободен (SIL OFL), кириллица
# родная; закрывает ГОСТ Р 7.0.97-2016 п. 3.3 о бесплатных шрифтах. Любая другая
# гарнитура в документе — нарушение: практика обязана выглядеть одинаково.
SPEC_FONT_BODY = "PT Serif"
SPEC_FONT_DISPLAY = SPEC_FONT_BODY
SPEC_FONT_MONO = SPEC_FONT_BODY
SPEC_FONT_TITLE = SPEC_FONT_BODY
SPEC_FONTS = {SPEC_FONT_BODY}
SPEC_FONT = SPEC_FONT_BODY  # обратная совместимость
SPEC_SIZES = {11.0, 12.0, 13.0, 14.0}
SPEC_LINE_SPACING = 1.15
SPEC_INDENT_CM = 1.25
TOLERANCE_MM = 1.0


def _fmt_margins(m: dict) -> str:
    return f"{m['top']}/{m['bottom']}/{m['left']}/{m['right']} мм"


MACHINE_RULES = (
    ("text.money", lambda: "каждая денежная сумма: цифры + пропись в круглых скобках; "
     "пропись сверяется с числом, prefix-валюты и общая валюта перечня тоже проверяются; "
     "сокращения тыс./млн./млрд. запрещены"),
    ("text.dates", lambda: "даты только ДД.ММ.ГГГГ"),
    ("text.yo", lambda: "буква ё/Ё запрещена"),
    ("text.quotes", lambda: "прямые кавычки запрещены, нужны елочки"),
    ("text.brackets", lambda: "квадратные скобки запрещены"),
    ("text.placeholders", lambda: "плейсхолдеры {{...}}, ХХХ/XXX и [ЗАПОЛНИТЬ...] запрещены"),
    ("text.blanks", lambda: "подчеркнутые пропуски ____ запрещены в процессуальных документах; "
     "в договорах под --dogovor допустимы"),
    ("docx.font", lambda: f"одна гарнитура: {SPEC_FONT_BODY}"),
    ("docx.size", lambda: f"кегли судебного профиля: {sorted(SPEC_SIZES)}; "
     f"адвокатский бланк: {sorted(SPEC_SIZES_ADVOKAT)}"),
    ("docx.style", lambda: "курсив, подчеркивание и гиперссылки запрещены; "
     "подчеркивание допустимо только в адвокатском бланке"),
    ("docx.paragraph", lambda: f"тело по ширине, интервал {SPEC_LINE_SPACING}x "
     f"({SPEC_LINE_SPACING_ADVOKAT}x для адвокатского бланка), отступ {SPEC_INDENT_CM} см "
     f"({sorted(SPEC_INDENT_CM_ADVOKAT)} для адвокатского бланка)"),
    ("docx.margins", lambda: f"поля: суд {_fmt_margins(SPEC_MARGINS_MM)}, "
     f"L3 {_fmt_margins(SPEC_MARGINS_MM_L3)}, договор {_fmt_margins(SPEC_MARGINS_MM_DOGOVOR)}, "
     f"адвокатский бланк {_fmt_margins(SPEC_MARGINS_MM_ADVOKAT)}"),
    ("docx.page", lambda: "номер страницы: настоящее поле PAGE в нижнем колонтитуле, по центру; "
     "первая страница без номера"),
    ("docx.attachments", lambda: "приложения: сквозная нумерация, каждое непроцессуальное приложение "
     "упомянуто в тексте"),
    ("docx.md-parity", lambda: ".md и .docx совпадают по числам и длинным абзацам при --md"),
)


def iter_paragraphs(doc):
    """Все абзацы документа: тело, таблицы (рекурсивно) и колонтитулы.

    `doc.paragraphs` возвращает ТОЛЬКО абзацы верхнего уровня. Шапка процессуального
    документа по нашему же стандарту — плавающая таблица, поэтому проверка шрифтов,
    кеглей и дат ее не видела, а сверка .md/.docx давала ложное «числа из .md
    отсутствуют в .docx»: числа лежали в ячейках. Найдено аудитом 03.08.2026.
    """
    def walk_container(container):
        for par in getattr(container, "paragraphs", []):
            yield par
        for table in getattr(container, "tables", []):
            for row in table.rows:
                for cell in row.cells:
                    yield from walk_container(cell)

    yield from walk_container(doc)
    for sec in doc.sections:
        for area in (sec.header, sec.footer, sec.first_page_header,
                     sec.first_page_footer, sec.even_page_header, sec.even_page_footer):
            try:
                yield from walk_container(area)
            except (AttributeError, ValueError):
                continue


def doc_text(doc) -> str:
    return "\n".join(p.text for p in iter_paragraphs(doc))


def _theme_fonts(doc) -> dict:
    """Гарнитуры темы документа (theme1.xml): major/minor → имя латиницы.

    Ран со ссылкой w:asciiTheme="majorHAnsi" наследует гарнитуру темы — в
    шаблоне python-docx это Cambria/Calibri, то есть чужая гарнитура при
    внешне зеленом документе (проба круга 6). Нет темы — None, не судим."""
    out = {"major": None, "minor": None}
    try:
        for part in doc.part.package.iter_parts():
            if str(part.partname).endswith("theme1.xml"):
                xml = part.blob.decode("utf-8", "replace")
                for key, tag in (("major", "majorFont"), ("minor", "minorFont")):
                    m = re.search(rf"<a:{tag}>\s*<a:latin[^>]*typeface=\"([^\"]*)\"",
                                  xml)
                    if m and m.group(1):
                        out[key] = m.group(1)
                break
    except Exception:
        pass
    return out


def _fonts_and_sizes(doc) -> tuple[set, set]:
    """Все гарнитуры и кегли, ДЕЙСТВУЮЩИЕ на документ: раны (document +
    колонтитулы, все атрибуты rFonts включая hAnsi — кириллицу Word берет
    оттуда), стили, реально примененные в документе (и Normal как база), и
    docDefaults. Тематические ссылки (asciiTheme/hAnsiTheme) судим только на
    ранах: в docDefaults шаблона они стоят по умолчанию и перекрыты везде —
    судить их там значит браковать каждый документ (ложная тревога)."""
    from docx.oxml.ns import qn

    fonts, sizes = set(), set()
    theme = _theme_fonts(doc)

    def collect_rpr(rpr, allow_theme):
        if rpr is None:
            return
        rf = rpr.find(qn("w:rFonts"))
        if rf is not None:
            for a in ("w:ascii", "w:hAnsi", "w:cs"):
                v = rf.get(qn(a))
                if v:
                    fonts.add(v)
            if allow_theme:
                for a in ("w:asciiTheme", "w:hAnsiTheme"):
                    v = rf.get(qn(a))
                    if not v:
                        continue
                    name = theme.get("major" if v.startswith("major") else "minor")
                    if name:
                        fonts.add(name)
        sz = rpr.find(qn("w:sz"))
        if sz is not None:
            v = sz.get(qn("w:val"))
            if v and v.isdigit():
                sizes.add(round(int(v) / 2, 1))

    parts = [doc.element]
    for sec in doc.sections:
        for area in (sec.header, sec.footer, sec.first_page_header,
                     sec.first_page_footer, sec.even_page_header, sec.even_page_footer):
            try:
                parts.append(area._element)
            except (AttributeError, ValueError):
                continue
    for el in parts:
        for r in el.iter(qn("w:r")):
            collect_rpr(r.find(qn("w:rPr")), allow_theme=True)

    styles_el = doc.styles.element
    by_id = {}
    for st in styles_el.findall(qn("w:style")):
        sid = st.get(qn("w:styleId"))
        if sid:
            by_id[sid] = st
    used = {"Normal"}
    for el in parts:
        for tag in ("w:pStyle", "w:tblStyle"):
            for ps in el.iter(qn(tag)):
                v = ps.get(qn("w:val"))
                if v:
                    used.add(v)
    seen, stack = set(), list(used)
    while stack:
        sid = stack.pop()
        if sid in seen:
            continue
        seen.add(sid)
        st = by_id.get(sid)
        if st is None:
            continue
        collect_rpr(st.find(qn("w:rPr")), allow_theme=False)
        base = st.find(qn("w:basedOn"))
        if base is not None and base.get(qn("w:val")):
            stack.append(base.get(qn("w:val")))

    dd = styles_el.find(qn("w:docDefaults"))
    if dd is not None:
        rprd = dd.find(qn("w:rPrDefault"))
        if rprd is not None:
            collect_rpr(rprd.find(qn("w:rPr")), allow_theme=False)
    return fonts, sizes


def _appendix_header(text: str) -> bool:
    return bool(re.fullmatch(r"\s*(?:Приложени[ея]|ПРИЛОЖЕНИ[ЕЯ])\s*:?\s*",
                             text or ""))


def _is_appendix_item(p) -> bool:
    """Пункт перечня приложений: номер ведет Word (w:numPr), в тексте абзаца
    его нет — по строке пункт не отличить, судим по структуре."""
    from docx.oxml.ns import qn
    ppr = p._p.find(qn("w:pPr"))
    return ppr is not None and ppr.find(qn("w:numPr")) is not None


def money_text(doc) -> str:
    """Текст для денежной проверки: тело без пунктов перечня приложений плюс
    таблицы, колонтитулы, поля форм (w:sdt) и надписи (w:txbxContent) целиком.

    Раньше срез приложений применялся к общей склейке, а `iter_paragraphs` отдает
    таблицы ПОСЛЕ абзацев верхнего уровня — поэтому строка «ПРИЛОЖЕНИЯ:» отрезала
    заодно все таблицы документа. Приложения есть в каждом иске, а расчет сумм
    и цена иска в шапке живут именно в таблицах: денежная проверка не работала
    ни в одном реальном иске (проба круга 5, 20.08.2026).

    Хвост за перечнем срезался целиком, а расчет цены иска — обязательная часть
    заявления и почти всегда стоит в самом конце, ПОСЛЕ перечня: ложная пропись
    там проходила незамеченной. Теперь пропускаются только сами пункты перечня
    (автонумерация w:numPr или ручная «1. »), текст после них проверяется
    (проба круга 6, 20.08.2026). Поля форм и надписи — обиход шаблонов Word и
    бланков: `doc.paragraphs` их не отдает, текст собирается из XML (та же
    проба).
    """
    telo_lines = []
    posle_zagolovka = False
    for p in getattr(doc, "paragraphs", []):
        if not posle_zagolovka and _appendix_header(p.text):
            posle_zagolovka = True
            continue
        if posle_zagolovka and (_is_appendix_item(p)
                                or re.match(r"^\s*\d{1,2}[.)]\s", p.text or "")):
            continue
        telo_lines.append(p.text)
    telo = "\n".join(telo_lines)
    prochee = []

    def _iz_yacheek(container):
        """Абзацы ячеек, включая вложенные таблицы. iter_paragraphs сюда не годится:
        она в конце просит doc.sections, которых у ячейки нет."""
        for par in getattr(container, "paragraphs", []):
            yield par
        for tbl in getattr(container, "tables", []):
            for r in tbl.rows:
                for c in r.cells:
                    yield from _iz_yacheek(c)

    # Между ячейками строки ставится непробельный барьер « | », как в разметке .md.
    # Раньше ячейки склеивались переводом строки, и `\s*` в денежной регулярке сшивал
    # конец одной ячейки с началом соседней: строка таблицы «… Фамилия И. Р. | 5 |
    # …» читалась как сумма «Р. 5» и требовала прописи у номера строки. Тот же текст
    # в .md сторож пропускал (там разделитель «|») — один документ давал два
    # противоположных ответа. Внутри ячейки перевод строки сохранен: сумма и пропись
    # часто стоят там разными абзацами, и барьер между ними дал бы обратную ложь.
    for table in getattr(doc, "tables", []):
        for row in table.rows:
            prochee.append(" | ".join(
                "\n".join(par.text for par in _iz_yacheek(cell))
                for cell in row.cells))
    # Поля форм (w:sdt) и надписи (w:txbxContent): doc.paragraphs их не видит,
    # а шаблоны Word и бланки набирают именно ими — ложная сумма внутри поля
    # или надписи проходила мимо всех проверок (проба круга 6, 20.08.2026).
    from docx.oxml.ns import qn
    for tag in ("w:sdtContent", "w:txbxContent"):
        for el in doc.element.body.iter(qn(tag)):
            txt = "".join(t.text or "" for t in el.iter(qn("w:t")))
            if txt.strip():
                prochee.append(txt)
    for sec in doc.sections:
        for part in (sec.header, sec.footer, getattr(sec, "first_page_header", None),
                     getattr(sec, "first_page_footer", None)):
            if part is not None:
                prochee.extend(par.text for par in getattr(part, "paragraphs", []))
    # Барьер « | » и МЕЖДУ элементами prochee, не только внутри строки таблицы:
    # «Р.» стояло последней ячейкой одной строки, а «5» — первой ячейкой следующей,
    # и перевод строки между ними денежная регулярка сшивала так же (проба 01.09.2026).
    return "\n".join([telo, " | ".join(prochee)])


def check_docx(path: str, l3: bool = False, dogovor: bool = False,
               advokat: bool = False) -> list[str]:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    problems: list[str] = []
    doc = Document(path)
    sec = doc.sections[0]
    spec = (SPEC_MARGINS_MM_ADVOKAT if advokat
            else SPEC_MARGINS_MM_DOGOVOR if dogovor
            else SPEC_MARGINS_MM_L3 if l3 else SPEC_MARGINS_MM)
    sizes_spec = SPEC_SIZES_ADVOKAT if advokat else SPEC_SIZES
    spacing_spec = SPEC_LINE_SPACING_ADVOKAT if advokat else SPEC_LINE_SPACING
    indent_spec = SPEC_INDENT_CM_ADVOKAT if advokat else {SPEC_INDENT_CM}
    got = {"top": sec.top_margin.mm, "bottom": sec.bottom_margin.mm,
           "left": sec.left_margin.mm, "right": sec.right_margin.mm}
    for side, want in spec.items():
        if abs(got[side] - want) > TOLERANCE_MM:
            problems.append(f"поле {side}: {got[side]:.0f} мм вместо {want} мм"
                            + (" (L3/кассация ВС)" if l3 and side == "left" else "")
                            + (" (бланк адвокатского центра)" if advokat else "")
                            + (" (профиль договора)" if dogovor and not advokat else ""))

    # Гарнитуру и кегль Word пишет НЕ ТОЛЬКО на ран: стиль абзаца (включая
    # Normal), docDefaults и тема документа задают их мимо ранов, а w:hAnsi —
    # тот атрибут, откуда Word берет КИРИЛЛИЦУ: ascii=PT Serif + hAnsi=Times
    # New Roman дает весь русский текст чужой гарнитурой при зеленом вердикте
    # (проба круга 6, 20.08.2026). Собираем объявления со ВСЕХ уровней,
    # которые реально действуют на этот документ.
    fonts, sizes = _fonts_and_sizes(doc)
    italic, underline = 0, 0
    for p in iter_paragraphs(doc):
        for r in p.runs:
            if r.font.italic:
                italic += 1
            if r.font.underline:
                underline += 1
    alien_fonts = fonts - SPEC_FONTS
    if alien_fonts:
        problems.append(f"чужие шрифты: {', '.join(sorted(alien_fonts))} "
                        f"(допустимы только {', '.join(sorted(SPEC_FONTS))})")
    alien_sizes = sizes - sizes_spec
    if alien_sizes:
        problems.append(f"кегли вне спецификации: {sorted(alien_sizes)} "
                        f"(допустимы {sorted(sizes_spec)})")
    if italic:
        problems.append(f"курсив в {italic} фрагментах — спецификация запрещает")
    # В бланке подчеркивание — это поля под заполнение («Тел. ______») и место
    # под подпись, а не оформление текста: запрет §3 писан для судебных документов.
    if underline and not advokat:
        problems.append(f"подчеркивание в {underline} фрагментах — спецификация запрещает")

    # Гиперссылка — отдельная ловушка: ее runs лежат внутри w:hyperlink, в
    # p.runs не попадают, и проверка выше их не видит. Рендерер оформляет их
    # стилем Hyperlink — синим с подчеркиванием — даже когда прямое
    # форматирование говорит обратное. Стандарт запрещает и то и другое
    # (DOCX_FORMATTING.md §3), поэтому ловим сам факт наличия гиперссылки
    # с видимым текстом (пропущено сторожем до 03.08.2026).
    doc_xml = doc.element.xml
    if "<w:hyperlink" in doc_xml:
        n = doc_xml.count("<w:hyperlink")
        problems.append(f"гиперссылок в тексте: {n} — рендерер красит их синим с "
                        "подчеркиванием поверх прямого форматирования; стандарт "
                        "запрещает цвет и подчеркивание. Навигацию давать обычным "
                        "текстом (add_static_toc с linked=False)")

    body = [p for p in doc.paragraphs if len(p.text.strip()) > 80]
    not_justified = [p for p in body
                     if p.paragraph_format.alignment not in (WD_ALIGN_PARAGRAPH.JUSTIFY, None)]
    if body and len(not_justified) > len(body) * 0.2:
        problems.append(f"тело не по ширине: {len(not_justified)} из {len(body)} длинных абзацев "
                        "выровнены иначе (нужен JUSTIFY)")

    spacings = {round(p.paragraph_format.line_spacing, 2) for p in body
                if p.paragraph_format.line_spacing}
    bad_spacing = spacings - {spacing_spec}
    if bad_spacing:
        problems.append(f"межстрочный интервал {sorted(bad_spacing)} вместо {spacing_spec}")

    # Отрицательный отступ первой строки — не ошибка, если это висячий отступ
    # нумерованного абзаца: left_indent положителен и гасит его ровно. Такой
    # абзац дает точную ссылку «пункт 14» и введен протоколом 03.08.2026.
    bad_indent = set()
    for par in body:
        fi = par.paragraph_format.first_line_indent
        if fi is None:
            continue
        fi_cm = round(fi.cm, 2)
        li = par.paragraph_format.left_indent
        li_cm = round(li.cm, 2) if li is not None else 0.0
        if fi_cm < 0 and abs(li_cm + fi_cm) <= 0.05 and li_cm > 0:
            continue  # корректный висячий отступ нумерованного абзаца
        if abs(fi_cm) <= 0.05 and li_cm > 0:
            continue  # блок-цитата: отступ слева есть, первой строки нет
        if all(abs(fi_cm - want) > 0.05 for want in indent_spec):
            bad_indent.add(fi_cm)
    if bad_indent:
        problems.append(f"абзацный отступ {sorted(bad_indent)} см вместо {sorted(indent_spec)} см "
                        "(висячий отступ нумерованного абзаца засчитывается, если "
                        "левый отступ гасит его ровно)")

    text = doc_text(doc)
    problems += check_text(text, os.path.basename(path), dogovor=dogovor or advokat)
    # Деньги проверяются по СТРУКТУРЕ документа, а не по склейке: перечень
    # приложений режется в теле, а таблицы и колонтитулы идут в проверку целиком.
    # Иначе строка «ПРИЛОЖЕНИЯ:» отрезала все таблицы разом (проба круга 5).
    problems = [p for p in problems if "пропис" not in p and "совпадает" not in p]
    problems += check_money_propis(money_text(doc), os.path.basename(path))

    # Нумерация страниц — безусловное требование протокола (решение владельца
    # 03.08.2026). Порога по объему больше нет: короткий документ тоже может
    # разойтись на две страницы после правки, а незамеченная потеря нумерации
    # обнаруживается уже в суде. DocBuilder.save() ставит поле сам — отсутствие
    # поля означает, что документ собран мимо DocBuilder.
    # Поле живет в колонтитуле — это ОТДЕЛЬНАЯ часть пакета .docx, и в
    # doc.element.xml его нет никогда. Прежняя проверка искала только там,
    # поэтому на корректном документе давала ложную тревогу, а на собранном
    # мимо DocBuilder — не давала никакой (баг найден 03.08.2026).
    def page_paragraphs(area):
        out = []
        try:
            for p in getattr(area, "paragraphs", []):
                xml = p._p.xml
                if re.search(r'<w:fldSimple[^>]+w:instr="[^"]*\bPAGE\b', xml) or "w:instrText" in xml and "PAGE" in xml:
                    out.append(p)
        except Exception:
            pass
        return out

    footer_pages, header_pages, first_footer_pages = [], [], []
    for sec in doc.sections:
        footer_pages += page_paragraphs(sec.footer) + page_paragraphs(sec.even_page_footer)
        first_footer_pages += page_paragraphs(sec.first_page_footer)
        header_pages += page_paragraphs(sec.header) + page_paragraphs(sec.first_page_header) + page_paragraphs(sec.even_page_header)
    if first_footer_pages:
        problems.append("номер страницы стоит на первой странице — первая страница должна быть без номера")
    if not footer_pages:
        if header_pages:
            # Место номера — тоже часть стандарта, и проверять его должна машина.
            problems.append("номер страницы стоит в ВЕРХНЕМ колонтитуле — с 04.08.2026 "
                            "он ставится в нижнем, по центру")
        else:
            problems.append("нет поля номера страницы — нумерация обязательна в каждом "
                            "документе без исключений (протокол 03.08.2026)")
    elif any(p.alignment != WD_ALIGN_PARAGRAPH.CENTER for p in footer_pages):
        problems.append("номер страницы в нижнем колонтитуле стоит не по центру")
    return problems


# Денежная сумма: число (пробелы/неразрывные или точка-между-тройками как
# разряды, запятая — копейки), затем опционально пропись в круглых скобках,
# затем валюта — либо валюта последним словом внутри скобок. Якорь — слово
# валюты: без него даты, статьи, номера дел, ИНН, ставки и листы дела сюда
# не попадают, и правильно — прописи требует только сумма денег.
# Разряд склеивается ТОЛЬКО пробелом/неразрывным/узким (Word и печатные формы
# ставят U+00A0, U+202F, U+2009 между тройками — без них проверка выключается
# пробелом, проба круга 4) или точкой между тройками:
# перевод строки числа не склеивает — иначе ячейка таблицы сливается с номером
# строки, а номер счета — с суммой под ним (ложные тревоги пробы 20.08.2026).
_MONEY_NUM = (r"\d{1,3}(?:[ \u00a0\u202f\u2009]\d{3})+(?:,\d{1,2})?"
              r"|\d{1,3}(?:\.\d{3})+(?:,\d{1,2})?"
              r"|\d{1,3}(?:,\d{3})+"
              r"|\d+(?:,\d{1,2})?")
# Словарь валют — ОДИН на проект, scripts/money_rule.py: рубли и копейки,
# доллары, евро, центы, символы и «р.». Сторож знал только рубли, вердикт —
# больше, и пропись в чужой валюте («200 000 (сто) долларов США», расхождение
# в 2000 раз) не сверялась вовсе (этап 9.22, круг 9). Склонение единицы —
# параметром валюты из того же словаря, не зашитое «рублей».
_MONEY_CUR = _mr.CUR_RE if _mr is not None else r"руб(?:л[а-я]*)?\.?|коп(?:е[а-я]*)?\.?|₽"
_MONEY_CUR_END = _mr.CUR_END if _mr is not None else r"(?![A-Za-zА-Яа-яЁё])"
_MONEY_RE = re.compile(
    rf"(?P<num>{_MONEY_NUM})"
    rf"(?:\s*\((?P<propis>[^()]*)\))?"
    rf"(?:\s+(?P<cur>{_MONEY_CUR}))?{_MONEY_CUR_END}",
    re.I)

# Сокращенная форма «500 тыс. руб.» / «12 млн руб.» — та же денежная сумма,
# и прописи она требует так же; само сокращение тысяч/миллионов в документе
# недопустимо (проба круга 4 этапа 9). Якорь-валюта обязателен: «5 тыс. штук»
# — не деньги.
_MONEY_ABBR_RE = re.compile(
    r"(?<![\dа-яА-Я])(?P<num>\d+(?:[ \u00a0\u202f\u2009.]\d{3})*(?:,\d+)?)"
    r"\s*(?P<scale>тыс|млн|млрд)\.?(?P<propis>\s*\([^()]*\))?"
    rf"\s+(?P<cur>{_MONEY_CUR}){_MONEY_CUR_END}",
    re.I)
_ABBR_MULT = {"тыс": 1_000, "млн": 1_000_000, "млрд": 1_000_000_000}
_MONEY_PREFIX_RE = re.compile(
    rf"(?P<cur>{_MONEY_CUR})\s*(?P<num>{_MONEY_NUM})"
    rf"(?:\s*\((?P<propis>[^()]*)\))?",
    re.I)


# Пропись ПЕРЕД числом — «двести тысяч (100 000) рублей» — та же
# контролирующая форма, только зеркальная: скобки тут несут цифры, а не слова.
# Без этой ветки число внутри скобок не деньги (закрывающая скобка отрезает
# его от слова валюты), и ложь о сумме проходила (проба круга 6, 20.08.2026).
_MONEY_BEFORE_PROPIS_RE = re.compile(
    rf"(?P<words>[а-яё][а-яе \-]{{1,90}}?)\(\s*(?P<num>{_MONEY_NUM})\s*\)"
    rf"\s*(?P<cur>{_MONEY_CUR}){_MONEY_CUR_END}",
    re.I)
# Слово-числительное — словарь словоформ propis.py из money_rule, а не
# префикс со свободным хвостом: прежний перечень добавил «тридцать/сорок/
# пятьдесят» и забыл «девяносто» — верная пропись «девяносто тысяч (90 000)
# рублей» браковалась (этап 9.22, круг 9); «полтора миллиона» сверяется
# структурно (money_rule.polutora_value).


def check_money_propis(text: str, where: str) -> list[str]:
    """Денежная сумма обязана нести пропись в круглых скобках, и пропись обязана
    совпадать с числом: «1 000 (сто тысяч) рублей» глазами не ловится, для того
    и прибор. Сверка — по словам целиком: «пять» внутри «пятьдесят» и префикс
    «одна тысяча» в «одна тысяча двести» совпадением не считаются. Разбор —
    единый, scripts/money_rule.py (словарь валют и числительных, цитаты,
    конвертер propis.py); недоступен — fail-closed. Число сверх предела
    конвертера — строка нарушения, а не трасса: до остальных проверок документа
    авария недопустима."""
    if _mr is None:
        return [f"{where}: scripts/money_rule.py (propis) недоступен — "
                f"совпадение прописи с числом НЕ проверено (fail-closed)"]
    # Перечень приложений — реквизиты прилагаемых документов, а не денежные
    # суммы документа: прописи он не несет никогда (ложная тревога круга 4).
    # Но текст ПОСЛЕ перечня — часть документа: расчет цены иска почти всегда
    # стоит в конце, и хвост, срезанный целиком, выпадал из проверки (проба
    # круга 6, 20.08.2026). Срезаются только пункты перечня (строки с номером
    # или маркером списка); первая не-пунктовая строка — продолжение документа
    # и проверяется. В .docx срез делает money_text по СТРУКТУРЕ (w:numPr),
    # здесь — плоский текст .md.
    m_app = re.search(r"(?im)^\s*(?:Приложени[ея]|ПРИЛОЖЕНИ[ЕЯ])\s*:?\s*$", text)
    if m_app:
        golova = text[:m_app.start()]
        hvost_lines = []
        spisok = True
        for line in text[m_app.end():].split("\n"):
            s = line.strip()
            if spisok and (not s or re.match(r"^(?:\d{1,2}[.)]|[—–\-•*])\s", s)):
                continue
            spisok = False
            hvost_lines.append(line)
        text = golova + "\n" + "\n".join(hvost_lines)
    # Дословная цитата нормы или судебного акта в кавычках-елочках
    # воспроизводится как в источнике — ТРЕБОВАТЬ там пропись значит запретить
    # цитирование (правило проекта — цитировать дословно; возражения и жалобы
    # почти всегда цитируют обжалуемый акт с суммой, этап 9.22, круг 9). Но
    # пропись, которая в цитате ЕСТЬ, обязана совпадать с числом: подмена
    # суммы внутри елочек — та же ложь (проба круга 6). Цитаты проверяются
    # отдельным проходом в режиме «не требовать, но сверять», а из основного
    # текста вырезаются, чтобы не задвоить находки.
    text, quotes = _mr.split_quotes(text)
    problems = []
    # Пропись ПЕРЕД числом: «двести тысяч (100 000) рублей». Хвост фразы из
    # числительных (словарь propis.py, не префикс) сверяется с числом в
    # скобках; совпала — внутреннее число второй раз не судим. Числительных
    # в хвосте нет — это не форма прописи (например, «расчет (100 000)
    # рублей») — число судит общий проход ниже.
    covered = []
    for m in _MONEY_PREFIX_RE.finditer(text):
        num_raw = m.group("num")
        words_raw = m.group("propis") or ""
        int_str, kop_str = _mr.money_int(num_raw)
        if not int_str or int(int_str) == 0:
            continue
        cur_name = _mr.currency_of(m.group("cur"))
        cur_info = _mr.CURRENCIES.get(cur_name) if cur_name else None
        cur_show = m.group("cur")
        gender = cur_info["gender"] if cur_info else "м"
        variants = _mr.propis_variants(int(int_str), gender=gender)
        words = re.sub(r"\s+", " ", words_raw.strip().lower()).split()
        if not words:
            problems.append(f"{where}: сумма {cur_show} {num_raw} без прописи в "
                            f"круглых скобках — денежная сумма пишется цифрами "
                            f"и прописью")
            covered.append(m.span())
            continue
        if variants and not _mr.words_match(words, variants):
            problems.append(f"{where}: пропись «{words_raw.strip()}» НЕ совпадает "
                            f"с числом {num_raw} — ожидалось "
                            f"«{_mr.propis_word(int(int_str), gender=gender)}»")
        covered.append(m.span())
    for m in _MONEY_BEFORE_PROPIS_RE.finditer(text):
        words_all = re.sub(r"\s+", " ", m.group("words").strip().lower()).split()
        tail = _mr.numeral_tail(words_all)
        if not tail:
            continue
        covered.append(m.span())
        int_str, _ = _mr.money_int(m.group("num"))
        if not int_str or int(int_str) == 0:
            continue
        n = int(int_str)
        if _mr.polutora_value(tail) == n:
            continue  # «полтора миллиона (1 500 000) рублей» — верная форма
        cur_name = _mr.currency_of(m.group("cur"))
        gender = _mr.CURRENCIES[cur_name]["gender"] if cur_name else "м"
        if not _mr.words_match(tail, _mr.propis_variants(n, gender=gender)):
            problems.append(f"{where}: пропись «{' '.join(tail)}» НЕ совпадает "
                            f"с числом {m.group('num')} — ожидалось "
                            f"«{_mr.propis_word(n, gender=gender)}»")
    for m in _MONEY_ABBR_RE.finditer(text):
        # «500 тыс. руб.» / «12 млн руб.» запрещены как сокращение суммы даже
        # при прописи: правила требуют полную цифровую запись.
        int_str, _ = _mr.money_int(m.group("num"))
        full = int(int_str) * _ABBR_MULT[m.group("scale").lower()]
        full_show = f"{full:,}".replace(",", " ")
        problems.append(f"{where}: сумма «{m.group(0).strip()}» сокращением "
                        f"тыс./млн. — недопустимо, денежная сумма пишется "
                        f"полностью цифрами и прописью: «{full_show} "
                        f"({_mr.propis_word(full)}) {m.group('cur')}»")
    for m in _MONEY_RE.finditer(text):
        if any(a <= m.start() < b for a, b in covered):
            continue  # число в скобках формы «пропись (число)» — уже судили
        num_raw = m.group("num")
        words_raw = m.group("propis")
        cur = m.group("cur")
        words = re.sub(r"\s+", " ", (words_raw or "").strip().lower()).split()
        words = [w for w in words if w]
        # Якорь-валюта: снаружи скобок или последним словом внутри них —
        # «1 000 (сто тысяч рублей)» та же денежная сумма. Хвост разбирается
        # СТРУКТУРНО: «… рублей 00 копеек» цифрами — тот же обиход, что
        # «ноль копеек» словом (этап 9.22, круг 9).
        has_cur = cur is not None
        inner_cur, words = _mr.strip_currency_tail(words)
        has_cur = has_cur or inner_cur
        if not has_cur:
            tail = text[m.end():m.end() + 120]
            shared = words_raw and re.search(rf"\bи\s+{_MONEY_NUM}[^.;\n]{{0,80}}{_MONEY_CUR}{_MONEY_CUR_END}", tail, re.I)
            if not shared:
                continue  # не деньги: дата, статья, номер дела, ИНН, ставка
            int_str, _ = _mr.money_int(num_raw)
            if int_str and int(int_str):
                variants = _mr.propis_variants(int(int_str), gender="м")
                if variants and not _mr.words_match(words, variants):
                    problems.append(f"{where}: пропись «{words_raw.strip()}» НЕ совпадает "
                                    f"с числом {num_raw} при общей валюте в перечне сумм")
            continue
        int_str, kop_str = _mr.money_int(num_raw)
        if not int_str or int(int_str) == 0:
            continue  # «рублей 00 копеек» — нулевые копейки цифрами это обиход
        # Валюта находки — каноническая, из единого словаря: род числительного
        # и склонение единицы параметром валюты («долларов», не «рублей»),
        # мелкая единица — ее строкой словаря (рубль → копейка ж.р.,
        # доллар/евро → цент м.р.). При отсутствии слова СНАРУЖИ скобок валюта
        # берется по ОСНОВНОЙ единице прописи (inner_cur из strip_currency_tail),
        # а НЕ по последнему слову хвоста копеек: иначе целое число согласуется
        # по роду копейки и «два рубля» ожидается как «две» (причина 4 ремонта
        # 25.08.2026 — форма арендодателя «…рубля девяносто одна копейка»).
        cur_name = _mr.currency_of(cur) if cur else inner_cur
        cur_info = _mr.CURRENCIES.get(cur_name) if cur_name else None
        # Показ валюты в подсказке: слово снаружи скобок как в тексте, иначе —
        # каноническая форма основной единицы («рублей»), не хвост копеек.
        cur_show = cur or (cur_info["forms"][2] if cur_info else "")
        gender = cur_info["gender"] if cur_info else "м"
        minor_info = (_mr.CURRENCIES[cur_info["minor"]]
                      if cur_info and cur_info.get("minor")
                      else _mr.CURRENCIES["копейка"])
        # Сверка принимает ЛЮБОЙ из шести падежей (просительная часть —
        # винительный: «взыскать одну тысячу»), но ложь ловится в каждом:
        # родительный от ДРУГОГО числа — брак (проба круга 6, 20.08.2026).
        exp_rub_variants = _mr.propis_variants(int(int_str), gender=gender)
        if not exp_rub_variants:
            problems.append(f"{where}: число {num_raw} не конвертируется — "
                            f"сумма осталась НЕ проверенной")
            continue
        expected = " ".join(exp_rub_variants[0])  # именительный — для подсказки
        # Подсказка обязана называть ПОЛНУЮ верную форму. Прежде она давала
        # пропись без копеек, а пропись в судебном документе — контролирующая
        # форма: исполнение подсказки меняло взыскиваемую сумму (проба круга 5).
        polnoe = expected
        if kop_str and int(kop_str):
            try:
                polnoe = (f"{expected} "
                          f"{_mr.sklonenie(int(int_str), cur_info['forms'] if cur_info else _mr.CURRENCIES['рубль']['forms'])} "
                          f"{_mr.propis_word(int(kop_str), gender=minor_info['gender'])} "
                          f"{_mr.sklonenie(int(kop_str), minor_info['forms'])}")
            except ValueError:
                pass
        if not words:
            problems.append(f"{where}: сумма {num_raw} {cur_show} без прописи в "
                            f"круглых скобках — денежная сумма пишется цифрами "
                            f"и прописью: «{num_raw} ({polnoe}) {cur_show}»")
            continue
        try:
            # Копейка — женский род: «одна копейка», «двадцать одна копейка» —
            # грамотный русский, а не несовпадение (ложная тревога круга 4).
            exp_kop_variants = (_mr.propis_variants(int(kop_str),
                                                    gender=minor_info["gender"])
                                if kop_str and int(kop_str) else None)
        except ValueError as e:
            problems.append(f"{where}: копейки {num_raw} не конвертируются ({e}) — "
                            f"сумма осталась НЕ проверенной")
            continue
        if exp_kop_variants is None:
            ok = _mr.words_match(words, exp_rub_variants)
        else:
            # Копейки: «одна тысяча двести тридцать четыре рубля пятьдесят
            # шесть копеек» — обе части внутри одних скобок, словами целиком,
            # каждая в своем падеже.
            ok = any(
                words[:len(v)] == v and len(words) > len(v)
                and words[len(v)].startswith(_mr.MAJOR_PREFIXES)
                and _mr.words_match(words[len(v) + 1:], exp_kop_variants)
                for v in exp_rub_variants)
        if not ok:
            problems.append(f"{where}: пропись «{words_raw.strip()}» НЕ совпадает "
                            f"с числом {num_raw} — ожидалось «{polnoe}»")
    for q in quotes:
        # Цитата в елочках: пропись не ТРЕБУЕТСЯ (дословная норма), но
        # присутствующая обязана совпадать с числом — подмена суммы внутри
        # кавычек та же ложь (проба круга 6, 20.08.2026).
        problems += _scan_quote(q, where)
    return problems


def _scan_quote(fragment: str, where: str) -> list[str]:
    """Проход по цитате в елочках: только СОВПАДЕНИЕ присутствующей прописи,
    без требования ее наличия (дословное цитирование нормы — правило проекта)."""
    out = []
    covered = []
    for m in _MONEY_BEFORE_PROPIS_RE.finditer(fragment):
        words_all = re.sub(r"\s+", " ", m.group("words").strip().lower()).split()
        tail = _mr.numeral_tail(words_all)
        if not tail:
            continue
        covered.append(m.span())
        int_str, _ = _mr.money_int(m.group("num"))
        if not int_str or int(int_str) == 0:
            continue
        n = int(int_str)
        if _mr.polutora_value(tail) == n:
            continue
        cur_name = _mr.currency_of(m.group("cur"))
        gender = _mr.CURRENCIES[cur_name]["gender"] if cur_name else "м"
        if not _mr.words_match(tail, _mr.propis_variants(n, gender=gender)):
            out.append(f"{where}: пропись «{' '.join(tail)}» в цитате НЕ совпадает "
                       f"с числом {m.group('num')} — цитата дословна, но подмена "
                       f"суммы в ней — та же ложь")
    for m in _MONEY_RE.finditer(fragment):
        if any(a <= m.start() < b for a, b in covered):
            continue
        words_raw = m.group("propis")
        cur = m.group("cur")
        words = re.sub(r"\s+", " ", (words_raw or "").strip().lower()).split()
        words = [w for w in words if w]
        if not words:
            continue  # прописи нет — в цитате это не нарушение
        has_cur = cur is not None
        inner_cur, words = _mr.strip_currency_tail(words)
        has_cur = has_cur or inner_cur
        if not has_cur or not words:
            continue
        int_str, kop_str = _mr.money_int(m.group("num"))
        if not int_str or int(int_str) == 0:
            continue
        # Валюта — по основной единице прописи (inner_cur), не по хвосту копеек
        # (причина 4 ремонта 25.08.2026).
        cur_name = _mr.currency_of(cur) if cur else inner_cur
        cur_info = _mr.CURRENCIES.get(cur_name) if cur_name else None
        gender = cur_info["gender"] if cur_info else "м"
        minor_info = (_mr.CURRENCIES[cur_info["minor"]]
                      if cur_info and cur_info.get("minor")
                      else _mr.CURRENCIES["копейка"])
        variants = _mr.propis_variants(int(int_str), gender=gender)
        if not variants:
            continue
        kop_variants = None
        if kop_str and int(kop_str):
            kop_variants = _mr.propis_variants(int(kop_str),
                                               gender=minor_info["gender"])
        if kop_variants is None:
            ok = _mr.words_match(words, variants)
        else:
            ok = any(
                words[:len(v)] == v and len(words) > len(v)
                and words[len(v)].startswith(_mr.MAJOR_PREFIXES)
                and _mr.words_match(words[len(v) + 1:], kop_variants)
                for v in variants)
        if not ok:
            out.append(f"{where}: пропись «{words_raw.strip()}» в цитате НЕ совпадает "
                       f"с числом {m.group('num')} — цитата дословна, но подмена "
                       f"суммы в ней — та же ложь")
    return out


def check_text(text: str, where: str, dogovor: bool = False) -> list[str]:
    problems = []
    # Двойная кодировка: UTF-8, прочитанный как latin-1 (unicode_escape и
    # прочие перекодировщики), — «Ð\x92Ð·Ñ\x8b...» вместо «Взыскать». Такой
    # документ поврежден по факту, а все проверки по тексту слепнут: якоря
    # слов (валюта, разделы) искажены до неузнаваемости (проба круга 4 —
    # сумма с узким пробелом U+202F уходила вместе с кодировкой).
    mojibake = len(re.findall(r"[ÐÑ][\x80-\xbf]", text))
    if mojibake >= 3:
        problems.append(f"{where}: текст поврежден двойной кодировкой (UTF-8, "
                        f"прочитанный как latin-1) — {mojibake} признаков; документ "
                        f"пересобрать из исходника, проверки по нему слепнут")
    yo = len(re.findall(r"[ёЁ]", text))
    if yo:
        problems.append(f"{where}: буква «ё» встречается {yo} раз — в документах проекта запрещена")
    bad_dates = re.findall(r"\b\d{4}-\d{2}-\d{2}\b|\b\d{1,2}/\d{1,2}/\d{2,4}\b", text)
    if bad_dates:
        problems.append(f"{where}: даты не в формате ДД.ММ.ГГГГ: {', '.join(sorted(set(bad_dates))[:5])}")
    placeholders = []
    for m in re.finditer(r"\{\{[^}]+\}\}|\bХХХ\b|\bXXX\b|\[ЗАПОЛНИТЬ[^\]]*\]", text):
        token = m.group(0)
        # «ХХХ» — действующая серия бланка полиса ОСАГО (РСА, с 2018), а не забытая
        # заглушка. Отличаем по контексту: «серия ХХХ № 0648968315». Ложная тревога
        # на подлинном реквизите приучает пролистывать вывод сторожа.
        if token in ("ХХХ", "XXX"):
            ctx = text[max(0, m.start() - 40):m.end() + 20]
            if re.search(r"сери\w*\s+(?:ХХХ|XXX)\b", ctx) or re.search(r"(?:ХХХ|XXX)\s*№\s*\d", ctx):
                continue
        placeholders.append(token)
    if placeholders:
        problems.append(f"{where}: незаполненные плейсхолдеры: {', '.join(sorted(set(placeholders))[:5])}")
    # Квадратные скобки: в российском судебном обиходе их не пишут — ни в ссылках
    # на нормы, ни в тексте. Только круглые. Решение владельца 10.08.2026.
    square = len(re.findall(r"[\[\]]", text))
    if square:
        problems.append(f"{where}: квадратные скобки — {square} шт. В документах "
                        "практики их не пишут, ссылки и вставки только в круглых")
    straight_quotes = text.count('"')
    if straight_quotes:
        problems.append(f"{where}: прямые кавычки — {straight_quotes} шт.; нужны елочки «...»")
    # Ряд подчеркиваний: в процессуальном документе — забытое поле, в договоре —
    # законное место подписи. Под флагом --dogovor это НЕ нарушение: прибор,
    # который сам пишет «в договоре норма» и тут же валит проверку, приучает
    # пролистывать его вывод, а вместе с ложной тревогой пролистают настоящую.
    blanks = len(re.findall(r"_{4,}", text))
    if blanks and not dogovor:
        problems.append(f"{where}: подчеркнутых пропусков {blanks} — в процессуальном "
                        "документе это незаполненное поле; если это договор, гнать с --dogovor")
    problems += check_money_propis(text, where)
    return problems


def docx_appendices(docx_path: str):
    """Раздел приложений из .docx: (текст до раздела, [(номер, заголовок)]).

    С 03.08.2026 номера пунктов ведет сам Word (w:numPr), и в тексте абзаца
    их больше НЕТ. Разбор по тексту здесь слеп — номера берутся из порядка
    нумерованных абзацев после заголовка «ПРИЛОЖЕНИЯ». Вернет пустой список,
    если нумерация ручная: тогда вызывающий разбирает текст, как раньше.
    """
    from docx import Document
    from docx.oxml.ns import qn

    doc = Document(docx_path)
    paragraphs = list(doc.paragraphs)
    head = None
    for i, par in enumerate(paragraphs):
        if re.fullmatch(r"\s*(?:Приложени[ея]|ПРИЛОЖЕНИ[ЕЯ])\s*:?\s*", par.text or ""):
            head = i
            break
    if head is None:
        return "\n".join(x.text for x in paragraphs), []

    items = []
    for par in paragraphs[head + 1:]:
        if par._p.find(qn("w:pPr")) is not None and \
                par._p.find(qn("w:pPr")).find(qn("w:numPr")) is not None:
            if par.text.strip():
                items.append((len(items) + 1, par.text.strip()))
    body = "\n".join(x.text for x in paragraphs[:head])
    return body, items


def _mentioned_nums(text):
    """Номера приложений, названные в тексте. Считает не только «Приложение 3»,
    но и перечисление «Приложения 3, 4» и диапазон «Приложения 1—4» с любым тире.
    Прежний разбор брал ТОЛЬКО первое число после слова и на документе с четырьмя
    однотипными приложениями требовал упомянуть те, что уже названы диапазоном
    (прогон 25.08.2026: текст говорил «Приложения 1—4», сторож требовал 2 и 4)."""
    found = set()
    for m in re.finditer(r"приложени[ияюе]\s*№?\s*([\d\s,;—–\-]{1,40})", text, re.I):
        chunk = m.group(1)
        # диапазон: 1—4 разворачиваем в 1,2,3,4
        for a, b in re.findall(r"(\d{1,2})\s*[-—–]\s*(\d{1,2})", chunk):
            a, b = int(a), int(b)
            if a <= b and b - a < 50:
                found.update(range(a, b + 1))
        chunk_wo_ranges = re.sub(r"\d{1,2}\s*[-—–]\s*\d{1,2}", " ", chunk)
        found.update(int(x) for x in re.findall(r"\d{1,2}", chunk_wo_ranges))
    return found


def check_attachments(text: str, items=None) -> list[str]:
    """Приложения: сквозная нумерация и упоминание каждого в тексте документа.

    items — готовый список (номер, заголовок) из автонумерации Word. Передан —
    нумерация заведомо сквозная (ее ведет программа), проверяется только
    упоминание в тексте.
    """
    if items:
        problems = []
        nums = [n for n, _ in items]
        PROCEDURAL = ("пошлин", "егрюл", "егрип", "направлени", "вручени",
                      "почтов", "квитанц", "опись", "доверенност", "диплом",
                      "ордер", "выписка из един")
        procedural_nums = {n for n, title in items
                           if any(k in title.lower() for k in PROCEDURAL)}
        mentioned = _mentioned_nums(text)
        missing = [n for n in nums
                   if n not in mentioned and n not in procedural_nums]
        if missing:
            problems.append(f"приложения {missing} не упомянуты в тексте документа")
        return problems

    m = re.search(r"(?im)^\s*(?:Приложени[ея]|ПРИЛОЖЕНИ[ЕЯ])\s*:?\s*$", text)
    if not m:
        return []
    tail = text[m.end():]
    body = text[:m.start()]
    items = re.findall(r"(?m)^\s*(\d{1,2})[.)]\s+(.+)$", tail)
    nums = [int(n) for n, _ in items]
    # Приложения по ст. 126 АПК и ст. 132 ГПК подаются в силу закона, а не в
    # подтверждение довода: госпошлина, выписка ЕГРЮЛ, доказательства
    # направления копии, доверенность, диплом. Требовать ссылку на них в
    # мотивировочной части бессмысленно — это была ложная тревога.
    PROCEDURAL = ("пошлин", "егрюл", "егрип", "направлени", "вручени",
                  "почтов", "квитанц", "опись", "доверенност", "диплом",
                  "ордер", "выписка из един")
    procedural_nums = {int(n) for n, title in items
                       if any(k in title.lower() for k in PROCEDURAL)}
    problems = []
    if not nums:
        return ["раздел «Приложения» есть, но перечень пуст или не пронумерован"]
    expected = list(range(1, len(nums) + 1))
    if nums != expected:
        problems.append(f"нумерация приложений не сквозная: {nums} (ожидалось {expected})")
    mentioned = _mentioned_nums(body)
    missing = [n for n in nums if n not in mentioned and n not in procedural_nums]
    if missing:
        problems.append(f"приложения {missing} не упомянуты в тексте документа")
    return problems


# Короткая строка (шапка, подпись, номер листа) совпадает по случайности и дает
# ложные срабатывания. Сверяем абзацы от этой длины — они уникальны.
MIN_PARA_FOR_DIFF = 60


def check_md_vs_docx(md_path: str, docx_path: str) -> list[str]:
    """Два файла — один документ. Разошлись — подадут не то, что проверяли."""
    from docx import Document

    def norm(s: str) -> str:
        s = re.sub(r"[*_`#>\|\-]+", " ", s)
        # Ведущий номер списка в .md набран текстом («1. приобщить…»), а в .docx
        # его рисует нумерация Word и в тексте абзаца его нет. Без этой стрижки
        # каждый пункт просительной части читался как расхождение файлов.
        s = re.sub(r"^\s*\d{1,3}\s*[.)]\s*", "", s)
        return re.sub(r"\s+", " ", s).strip().lower()

    md_raw = open(md_path, encoding="utf-8", errors="replace").read()
    md = norm(md_raw)
    dx = norm(doc_text(Document(docx_path)))
    problems = []
    ratio = len(dx) / max(len(md), 1)
    if not 0.7 <= ratio <= 1.4:
        problems.append(f"объем .docx отличается от .md в {ratio:.2f} раза "
                        f"({len(dx)} против {len(md)} знаков) — файлы разошлись")
    # ключевые числа: сумма из .md обязана быть в .docx
    md_nums = set(re.findall(r"\b\d[\d\s]{4,}\b", md))
    lost = [n for n in md_nums if re.sub(r"\s", "", n) not in re.sub(r"\s", "", dx)]
    if lost:
        problems.append(f"числа из .md отсутствуют в .docx: {', '.join(sorted(lost)[:5])}")

    # Объем и числа рассогласование ТЕКСТА не ловят. Прецедент 21.08.2026: .md
    # правился по замечаниям рецензента, .docx остался прежним — четыре внесенных
    # правки в собранный документ не попали, а сторож дал код 0 и «согласованность
    # в порядке». Объем изменился на доли процента, крупные числа не тронуты —
    # обе прежние проверки прошли, и владельцу ушла бы редакция, которую рецензент
    # завернул. Поэтому сверяем сам текст: каждый содержательный абзац .docx обязан
    # находиться в .md. Абзац из .docx, которого в .md нет, — это либо непересобранный
    # файл, либо вставка мимо одобренной редакции; оба случая означают «подадут не то,
    # что проверяли».
    # Сверяем НАПРАВЛЕННО: .md → .docx, а не наоборот. Обратное направление
    # ссорит два прибора: `create_docx._matches_approved` намеренно исключает
    # шапку и подпись из сверки с одобренной редакцией, потому что их в .md нет,
    # — а сторож, требующий найти в .md каждый абзац .docx, ровно на них и падал
    # (21.08.2026, четыре лишних круга к рецензенту, составитель обошел сторожа
    # короткой шапкой). Правка же, ради которой проверка и заводилась, живет
    # в .md и обязана доехать до .docx — это направление ее и ловит, а служебные
    # элементы сборщика ему не мешают.
    for para in md_raw.split("\n"):
        piece = norm(para)
        if len(piece) < MIN_PARA_FOR_DIFF:
            continue
        if piece not in dx:
            problems.append(
                "текст .md не доехал до .docx — абзац есть в .md, но отсутствует "
                f"в .docx: «{piece[:90]}…». Пересобрать .docx из текущего .md")
            break
    return problems


def _add_page_field(doc, top=False, center=True, first=False):
    """Поле PAGE в нижнем колонтитуле — как это делает DocBuilder.

    top=True — заведомо неверное место (верх): фикстура для проверки, что
    сторож это ловит.
    """
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    sec = doc.sections[0]
    if first:
        sec.different_first_page_header_footer = True
    area = sec.first_page_footer if first else (sec.header if top else sec.footer)
    area.is_linked_to_previous = False
    p = area.paragraphs[0] if area.paragraphs else area.add_paragraph()
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), " PAGE ")
    p._p.append(fld)
    return p


def _add_fake_page_text(doc):
    p = doc.sections[0].footer.paragraphs[0]
    p.text = "PAGE"
    return p


def _add_fake_hyperlink(doc):
    """Гиперссылка как ее ставит Word: runs внутри w:hyperlink, в p.runs не видны."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    p = doc.add_paragraph()
    link = OxmlElement("w:hyperlink")
    link.set(qn("w:anchor"), "toc1")
    run = OxmlElement("w:r")
    text = OxmlElement("w:t")
    text.text = "к разделу 1"
    run.append(text)
    link.append(run)
    p._p.append(link)
    return p


def selftest() -> int:
    import tempfile
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Mm, Pt, Cm

    tmp = tempfile.mkdtemp()

    def build(path, *, font=SPEC_FONT_BODY, size=12, left=30, top=20, bottom=30,
              right=15, italic=False, underline=False, spacing=1.15,
              indent=1.25, pages=True, pages_top=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
              table=None, hyperlink=False, paragraphs=1,
              text="Текст документа, достаточно длинный абзац для проверки выравнивания.",
              fake_page_text=False, page_center=True, page_first=False):
        """Фикстура документа. КАЖДЫЙ параметр обязан быть покрыт проверкой.

        Прежняя версия объявляла spacing= и indent=, но ни одна фикстура их не
        передавала, таблиц не строила вовсе (`grep -c add_table` → 0), и восемь
        ветвей сторожа можно было молча удалить при зеленом selftest (аудит
        03.08.2026). Шапка процессуального документа — плавающая таблица, то есть
        непокрытой оставалась ровно та часть, ради которой сторож и написан.
        """
        d = Document()
        sec = d.sections[0]
        sec.top_margin, sec.bottom_margin = Mm(top), Mm(bottom)
        sec.left_margin, sec.right_margin = Mm(left), Mm(right)
        if table is not None:
            t = d.add_table(rows=len(table), cols=len(table[0]))
            for row, cells in zip(t.rows, table):
                for cell, spec in zip(row.cells, cells):
                    cp = cell.paragraphs[0]
                    cr = cp.add_run(spec["text"])
                    cr.font.name = spec.get("font", font)
                    cr.font.size = Pt(spec.get("size", size))
                    if spec.get("underline"):
                        cr.font.underline = True
        for _ in range(paragraphs):
            p = d.add_paragraph()
            p.alignment = align
            p.paragraph_format.line_spacing = spacing
            p.paragraph_format.first_line_indent = Cm(indent)
            r = p.add_run(text * 2)
            r.font.name, r.font.size = font, Pt(size)
            r.font.italic, r.font.underline = italic, underline
        if hyperlink:
            _add_fake_hyperlink(d)
        if fake_page_text:
            _add_fake_page_text(d)
        elif pages:
            _add_page_field(d, top=pages_top, center=page_center, first=page_first)
        d.save(path)
        return path

    good = build(os.path.join(tmp, "good.docx"))
    no_pages = build(os.path.join(tmp, "nopage.docx"), pages=False)
    pages_up = build(os.path.join(tmp, "pageup.docx"), pages_top=True)
    page_fake = build(os.path.join(tmp, "pagefake.docx"), pages=False, fake_page_text=True)
    page_left = build(os.path.join(tmp, "pageleft.docx"), page_center=False)
    page_first = build(os.path.join(tmp, "pagefirst.docx"), page_first=True)
    bad_font = build(os.path.join(tmp, "font.docx"), font="Arial")
    bad_margin = build(os.path.join(tmp, "margin.docx"), left=25)
    l3_ok = build(os.path.join(tmp, "l3.docx"), left=35)
    with_italic = build(os.path.join(tmp, "it.docx"), italic=True)
    with_yo = build(os.path.join(tmp, "yo.docx"), text="Ёжик всё ещё идёт в суд. ")

    md_ok = os.path.join(tmp, "ok.md")
    open(md_ok, "w", encoding="utf-8").write(
        "Текст документа, достаточно длинный абзац для проверки выравнивания." * 2)
    md_other = os.path.join(tmp, "other.md")
    open(md_other, "w", encoding="utf-8").write("Совсем другой короткий текст. Сумма 1 250 000 руб.")

    att_ok = ("Прошу приобщить приложение 1 и приложение 2.\n\nПриложения:\n"
              "1. Договор\n2. Квитанция\n")
    att_gap = "Прошу приобщить приложение 1.\n\nПриложения:\n1. Договор\n3. Квитанция\n"
    att_none = "Исковое заявление.\n\nПриложения:\n1. Договор\n2. Акт сверки\n"
    txt_file = os.path.join(tmp, "not-md.txt")
    open(txt_file, "w", encoding="utf-8").write("текст")

    # Шапка процессуального документа — плавающая таблица. До 04.08.2026 ни одна
    # фикстура таблиц не строила, и весь обход iter_paragraphs был не покрыт.
    head_ok = [[{"text": "В Вахитовский районный суд города Казани"}],
               [{"text": "Истец: Иванов Иван Иванович"}]]
    head_alien_font = [[{"text": "В Вахитовский районный суд",
                         "font": "Times New Roman"}]]
    head_alien_size = [[{"text": "В Вахитовский районный суд", "size": 9}]]
    head_underline = [[{"text": "В Вахитовский районный суд", "underline": True}]]
    head_numbers = [[{"text": "Цена иска: 1 250 000 руб."}]]

    with_table = build(os.path.join(tmp, "table.docx"), table=head_ok)
    tbl_font = build(os.path.join(tmp, "tblfont.docx"), table=head_alien_font)
    tbl_size = build(os.path.join(tmp, "tblsize.docx"), table=head_alien_size)
    tbl_underline = build(os.path.join(tmp, "tblund.docx"), table=head_underline)
    tbl_numbers = build(os.path.join(tmp, "tblnum.docx"), table=head_numbers)
    bad_spacing = build(os.path.join(tmp, "spacing.docx"), spacing=2.0)
    bad_indent = build(os.path.join(tmp, "indent.docx"), indent=0.5)
    bad_size = build(os.path.join(tmp, "size.docx"), size=9)
    with_underline = build(os.path.join(tmp, "und.docx"), underline=True)
    with_link = build(os.path.join(tmp, "link.docx"), hyperlink=True)
    left_aligned = build(os.path.join(tmp, "left.docx"),
                         align=WD_ALIGN_PARAGRAPH.LEFT, paragraphs=5)
    mostly_justified = build(os.path.join(tmp, "mostly.docx"), paragraphs=5)
    dogovor_ok = build(os.path.join(tmp, "dog.docx"), top=20, bottom=20, left=25, right=20)

    md_table = os.path.join(tmp, "table.md")
    open(md_table, "w", encoding="utf-8").write(
        "Цена иска: 1 250 000 руб.\n\n"
        + "Текст документа, достаточно длинный абзац для проверки выравнивания." * 2)

    checks = [
        ("корректный документ проходит", check_docx(good) == []),
        # Обход таблиц. Каждая ветвь проверки обязана видеть содержимое ячейки —
        # иначе шапка документа не проверяется вообще ничем.
        ("документ с таблицей-шапкой проходит", check_docx(with_table) == []),
        ("чужой шрифт В ЯЧЕЙКЕ шапки пойман",
         any("шрифт" in p for p in check_docx(tbl_font))),
        ("чужой кегль В ЯЧЕЙКЕ шапки пойман",
         any("кегли" in p for p in check_docx(tbl_size))),
        ("подчеркивание В ЯЧЕЙКЕ шапки поймано",
         any("подчеркивание" in p for p in check_docx(tbl_underline))),
        ("число из .md найдено в ЯЧЕЙКЕ таблицы, а не объявлено пропавшим",
         check_md_vs_docx(md_table, tbl_numbers) == []),
        # Интервал и отступ: параметры build() существовали, но не передавались.
        ("межстрочный интервал 2.0 пойман",
         any("интервал" in p for p in check_docx(bad_spacing))),
        ("абзацный отступ 0.5 см пойман",
         any("отступ" in p for p in check_docx(bad_indent))),
        ("кегль 9 пт пойман", any("кегли" in p for p in check_docx(bad_size))),
        ("подчеркивание в теле поймано",
         any("подчеркивание" in p for p in check_docx(with_underline))),
        ("гиперссылка поймана", any("гиперссыл" in p for p in check_docx(with_link))),
        # Порог JUSTIFY: доля, а не факт. Один абзац из пяти — терпимо, все пять — нет.
        ("сплошное выравнивание влево поймано",
         any("не по ширине" in p for p in check_docx(left_aligned))),
        ("документ по ширине порога не превышает",
         not any("не по ширине" in p for p in check_docx(mostly_justified))),
        # Профиль договора: свои поля 20/20/25/20.
        ("договор со своими полями проходит при --dogovor",
         check_docx(dogovor_ok, dogovor=True) == []),
        ("поля договора без флага считаются нарушением",
         any("поле" in p for p in check_docx(dogovor_ok))),
        ("судебные поля под флагом договора считаются нарушением",
         any("поле" in p for p in check_docx(good, dogovor=True))),
        ("чужой шрифт пойман", any("шрифт" in p for p in check_docx(bad_font))),
        ("поле не по спецификации поймано", any("поле left" in p for p in check_docx(bad_margin))),
        ("L3 с полем 35 проходит при --l3", check_docx(l3_ok, l3=True) == []),
        ("L3-поле без флага считается ошибкой", any("поле left" in p for p in check_docx(l3_ok))),
        ("курсив пойман", any("курсив" in p for p in check_docx(with_italic))),
        ("буква ё поймана", any("«ё»" in p for p in check_docx(with_yo))),
        ("отсутствие нумерации страниц поймано",
         any("номера страницы" in p for p in check_docx(no_pages))),
        ("обычный текст PAGE вместо поля пойман",
         any("номера страницы" in p for p in check_docx(page_fake))),
        # Место номера — часть стандарта: с 04.08.2026 он внизу по центру.
        ("номер страницы наверху пойман",
         any("ВЕРХНЕМ" in p for p in check_docx(pages_up))),
        ("номер страницы не по центру пойман",
         any("не по центру" in p for p in check_docx(page_left))),
        ("номер на первой странице пойман",
         any("первой странице" in p for p in check_docx(page_first))),
        ("номер страницы внизу претензий не вызывает",
         not any("номер страницы" in p for p in check_docx(good))),
        ("совпадающие md и docx проходят", check_md_vs_docx(md_ok, good) == []),
        ("разошедшиеся md и docx пойманы", check_md_vs_docx(md_other, good) != []),
        ("сквозная нумерация приложений проходит", check_attachments(att_ok) == []),
        ("дыра в нумерации приложений поймана", any("сквозная" in p for p in check_attachments(att_gap))),
        ("ноль упоминаний приложений пойман",
         any("не упомянуты" in p for p in check_attachments(att_none))),
        ("плейсхолдер пойман", any("плейсхолдер" in p for p in check_text("Сумма {{amount}}", "t"))),
        ("голое ХХХ поймано", any("плейсхолдер" in p for p in check_text("Полис ХХХ выдан", "t"))),
        ("серия полиса ОСАГО претензий не вызывает",
         check_text("полис ОСАГО серия ХХХ № 0648968315", "t") == []),
        ("дата не того формата поймана", any("ДД.ММ.ГГГГ" in p for p in check_text("2026-08-03", "t"))),
        # Ложная тревога дороже молчания: она приучает пролистывать вывод.
        ("подчеркнутый пропуск в процессуальном документе пойман",
         any("незаполненное поле" in p for p in check_text("Подпись ________", "t"))),
        ("в договоре место подписи претензий не вызывает",
         check_text("Подпись ________", "t", dogovor=True) == []),
        ("квадратные скобки пойманы",
         any("квадратные скобки" in p for p in check_text("(ст. 617 ГК РФ) [ст. 621 ГК РФ]", "t"))),
        ("круглые скобки претензий не вызывают",
         check_text("(ст. 617 ГК РФ)", "t") == []),
        ("прямые кавычки пойманы",
         any("прямые кавычки" in p for p in check_text('"Ромашка"', "t"))),
        # Пропись денежной суммы (этап 9.8): обе оси — сумма без прописи и
        # несовпадение ловятся, верная пропись и обиход молчат.
        ("сумма без прописи поймана",
         any("без прописи" in p for p in check_text(
             "Прошу взыскать с ответчика 100 000 рублей неустойки (ст. 330 ГК РФ).", "t"))),
        ("несовпадающая пропись поймана — «1 000 (сто тысяч)» глазами не ловится",
         any("НЕ совпадает" in p for p in check_text(
             "Прошу взыскать 1 000 (сто тысяч) рублей неустойки.", "t"))),
        ("верная пропись проходит",
         check_text("Прошу взыскать 1 000 (одна тысяча) рублей неустойки "
                    "(ст. 330 ГК РФ).", "t") == []),
        ("сокращение суммы запрещено даже с прописью",
         any("сокращением" in p for p in check_text(
             "Прошу взыскать 500 тыс. (пятьсот тысяч) руб.", "t"))),
        ("prefix-валюта без прописи поймана",
         any("без прописи" in p for p in check_text("Прошу взыскать $ 100 000.", "t"))),
        ("общая валюта для нескольких сумм сверяет первую пропись",
         any("общей валюте" in p for p in check_text(
             "Прошу взыскать 100 000 (девятьсот) и 50 000 (пятьдесят тысяч) рублей.", "t"))),
        ("пропись с копейками проходит",
         not any("пропис" in p or "совпадает" in p for p in check_text(
             "Прошу взыскать 1 234,56 (одна тысяча двести тридцать четыре рубля "
             "пятьдесят шесть копеек).", "t"))),
        ("нулевые копейки цифрами — обиход, не нарушение",
         not any("пропис" in p for p in check_text(
             "Прошу взыскать 5 000 (пять тысяч) рублей 00 копеек.", "t"))),
        # Этап 9.22, круг 9: обе оси денежного правила после сведения копий
        # в scripts/money_rule.py.
        ("нулевые копейки цифрами ВНУТРИ скобок — стандартная форма, не брак",
         check_text("Прошу взыскать задолженность 50 000,00 (пятьдесят тысяч "
                    "рублей 00 копеек) по договору (ст. 309 ГК РФ).", "t") == []),
        ("«девяносто тысяч (90 000)» — верная пропись, а не несовпадение",
         check_text("Прошу взыскать девяносто тысяч (90 000) рублей "
                    "задолженности.", "t") == []),
        ("«полтора миллиона (1 500 000)» — верная живая форма",
         check_text("Прошу взыскать полтора миллиона (1 500 000) рублей "
                    "задолженности.", "t") == []),
        ("пропись в чужой валюте сверяется: «200 000 (сто) долларов США» — брак",
         any("НЕ совпадает" in p for p in check_text(
             "Прошу взыскать 200 000 (сто) долларов США по контракту "
             "(ст. 317 ГК РФ).", "t"))),
        ("верная пропись в чужой валюте проходит",
         check_text("Прошу взыскать 200 000 (двести тысяч) долларов США "
                    "по контракту (ст. 317 ГК РФ).", "t") == []),
        # Причина 4 ремонта 25.08.2026: сумма с копейками на «1» — валюта
        # определялась по хвосту «копейка» (ж.р.), и целое «два» ждалось как
        # «две». Классическая форма арендодателя обязана проходить в обоих видах.
        ("копейки на «1»: классическая форма (валюта из прописи) проходит",
         check_text("Прошу взыскать задолженность 356 462,91 (триста пятьдесят "
                    "шесть тысяч четыреста шестьдесят два рубля девяносто одна "
                    "копейка) по договору аренды.", "t") == []),
        ("копейки на «1»: та же сумма с валютным словом после скобок проходит",
         check_text("Прошу взыскать 356 462,91 (триста пятьдесят шесть тысяч "
                    "четыреста шестьдесят два рубля девяносто одна копейка) руб.",
                    "t") == []),
        ("копейки на «1»: подмена целого в такой форме ловится",
         any("НЕ совпадает" in p for p in check_text(
             "Прошу взыскать 356 462,91 (триста пятьдесят шесть тысяч "
             "четыреста шестьдесят три рубля девяносто одна копейка).", "t"))),
        ("даты, статьи, номера дел, ИНН, ставки и листы прописи НЕ требуют",
         check_text("Заседание назначено на 21.08.2026 (ст. 333 ГК РФ, п. 71) "
                    "по делу № А65-123/2026, ИНН 1655021805, ставка 7,5 % "
                    "годовых, лист дела 82.", "t") == []),
        ("--md-only отвергает не .md",
         __import__("subprocess").run([sys.executable, __file__, "--md-only", txt_file],
                                      stdout=__import__("subprocess").DEVNULL,
                                      stderr=__import__("subprocess").DEVNULL).returncode == 2),
        ("--rules строится из MACHINE_RULES",
         all(rid in rules_report() for rid, _ in MACHINE_RULES)),
    ]
    for name, ok in checks:
        print(f"  {'✓' if ok else '✗'} {name}")
    bad = [n for n, ok in checks if not ok]
    print(f"selftest {'пройден' if not bad else 'ПРОВАЛЕН'}: {len(checks) - len(bad)}/{len(checks)}")
    return 1 if bad else 0


def rules_report() -> str:
    """Свод требований строится из MACHINE_RULES, а не из второго текста."""
    lines = [
        "СВОД МАШИННЫХ ТРЕБОВАНИЙ К ДОКУМЕНТУ (document_guard.py)",
        "Источник правил — MACHINE_RULES + константы SPEC_* этого же прибора.",
        "",
    ]
    lines += [f"• {rid}: {describe()}" for rid, describe in MACHINE_RULES]
    return "\n".join(lines)


def _source_line(path: str, units: int, unit_word: str) -> str:
    """Шапка вердикта: что именно прочитано. Вердикт без указания источника
    проверки — не вердикт (Часть 9 ремонта 25.08.2026)."""
    try:
        size = os.path.getsize(path)
    except OSError:
        size = 0
    return f"Проверено: {path} · {size} байт · {units} {unit_word}"


def main() -> int:
    ap = argparse.ArgumentParser(description="Проверка формата .docx и согласованности с .md")
    ap.add_argument("docx", nargs="?", help="путь к .docx")
    ap.add_argument("--md", help="парный .md для сверки")
    ap.add_argument("--l3", action="store_true", help="L3 или кассация ВС: левое поле 35 мм")
    ap.add_argument("--dogovor", action="store_true",
                    help="договор и прочее несудебное: поля 20/20/25/20 мм")
    ap.add_argument("--dogovor-advokat", dest="advokat", action="store_true",
                    help="бланк договора адвокатского центра (DOCX_FORMATTING.md §8): "
                         "поля 20/20/15/10 мм, кегли бланка, подчеркивание допустимо")
    ap.add_argument("--selftest", action="store_true")
    ap.add_argument("--md-only", dest="md_only", metavar="ЧЕРНОВИК.md",
                    help="проверить один .md без .docx (пропись, даты, «ё», скобки, "
                         "пропуски) — размыкает круг «формат требует .docx» до сборки")
    ap.add_argument("--rules", action="store_true",
                    help="напечатать полный свод проверяемых требований и выйти")
    a = ap.parse_args()
    if a.selftest:
        return selftest()
    if a.rules:
        print(rules_report())
        return 0
    if a.md_only:
        if not os.path.isfile(a.md_only):
            print(f"нет файла {a.md_only}", file=sys.stderr)
            return 2
        if os.path.splitext(a.md_only)[1].lower() != ".md":
            print(f"{a.md_only}: --md-only принимает только .md", file=sys.stderr)
            return 2
        text = open(a.md_only, encoding="utf-8").read()
        units = sum(1 for ln in text.split("\n") if ln.strip())
        print(_source_line(a.md_only, units, "непустых строк"))
        problems = check_text(text, os.path.basename(a.md_only),
                              dogovor=a.dogovor or a.advokat)
        if not problems:
            print(f"✓ {os.path.basename(a.md_only)}: текст в порядке")
            return 0
        print(f"⚠ {os.path.basename(a.md_only)}: нарушений {len(problems)}")
        for p in problems:
            print(f"   • {p}")
        return 1
    if not a.docx:
        ap.print_help()
        return 2
    if os.path.basename(a.docx).startswith("~$"):
        print(f"{a.docx} — временный файл Word, а не документ. Проверять нечего.", file=sys.stderr)
        return 2
    if not os.path.isfile(a.docx):
        print(f"нет файла {a.docx}", file=sys.stderr)
        return 2

    from docx import Document as _Doc
    print(_source_line(a.docx, len(list(iter_paragraphs(_Doc(a.docx)))), "абзацев"))
    problems = check_docx(a.docx, a.l3, a.dogovor, a.advokat)
    body, items = docx_appendices(a.docx)
    if items:
        problems += check_attachments(body, items)
    else:
        from docx import Document
        text = "\n".join(p.text for p in Document(a.docx).paragraphs)
        problems += check_attachments(text)
    if a.md:
        problems += check_md_vs_docx(a.md, a.docx)
        # .md уходит доверителю тем же документом, поэтому проверяется ЦЕЛИКОМ
        # той же машиной: плейсхолдеры, «ё», даты, скобки, пропуски и денежная
        # ось — иначе незаполненное поле переживает пересборку и всплывает в
        # следующей редакции (проба 20.08.2026).
        md_text = open(a.md, encoding="utf-8").read()
        problems += check_text(md_text, os.path.basename(a.md), dogovor=a.dogovor or a.advokat)

    if not problems:
        print(f"✓ {os.path.basename(a.docx)}: формат и согласованность в порядке")
        return 0
    print(f"⚠ {os.path.basename(a.docx)}: нарушений {len(problems)}")
    for p in problems:
        print(f"   • {p}")
    return 1


if __name__ == "__main__":
    sys.exit(main())
