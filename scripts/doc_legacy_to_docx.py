#!/usr/bin/env python3
"""Пересборка договора из HTML-слепка оригинала (.doc -> textutil html) в .docx.

Оригинал — Word 97 (.doc), редактировать его на месте нечем: LibreOffice/Word
недоступны. textutil сохраняет ВСЮ верстку в html (таблицы, кегли, выравнивания,
границы ячеек), поэтому html и есть носитель формата, а docx собирается по нему
один в один. Меняется только содержимое реквизитов и текст Заявки.
"""
import datetime
import html as htmllib
import os
import re
import shutil
import subprocess
import sys
import tempfile
from html.parser import HTMLParser

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor

SRC_HTML = os.environ.get("LEGACY_DOC_HTML", "/tmp/vp2026/src.html")

# --- геометрия страницы ------------------------------------------------------
# Оригинал — бланк с текстовой полосой 524.5 pt (185 мм). Точные поля из .doc не
# извлекаются, но полосу держим ту же: иначе таблицы сжимаются и текст в шапках
# начинает переноситься. A4, слева шире под подшивку.
PAGE_W, PAGE_H = 210.0, 297.0
M_TOP, M_BOTTOM, M_LEFT, M_RIGHT = 20.0, 20.0, 15.0, 10.0
TEXT_W_MM = PAGE_W - M_LEFT - M_RIGHT          # 165 мм
SRC_TEXT_W_PT = 524.5                          # ширина полосы в оригинале
PT_TO_MM = 25.4 / 72.0
SCALE = TEXT_W_MM / (SRC_TEXT_W_PT * PT_TO_MM)

# Правка владельца 20.08.2026 на договоре оказания услуг адвокатского центра:
# весь документ — PT Serif (гарнитура практики), таблицы-перечни ужимаются до 9 pt,
# колонтитул целиком 8 pt. Бланк наследует эти значения, а не Times New Roman образца.
FONT_FAMILY = "PT Serif"
TABLE_SIZE_MAP = {9.5: 9.0}
FOOTER_SIZE = 8.0

ALIGN = {
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
    "left": WD_ALIGN_PARAGRAPH.LEFT,
}


def parse_css(src: str) -> tuple[dict, dict]:
    """p.pN и td.tdN -> словари свойств."""
    p_styles, td_styles = {}, {}
    for name, body in re.findall(r"(p\.p\d+|td\.td\d+)\s*\{([^}]*)\}", src):
        props = {}
        for decl in body.split(";"):
            if ":" not in decl:
                continue
            k, v = decl.split(":", 1)
            props[k.strip()] = v.strip()
        key = name.split(".")[1]
        (p_styles if name.startswith("p.") else td_styles)[key] = props
    return p_styles, td_styles


def font_of(props: dict) -> tuple[float, str]:
    """'font: 9.5px Times New Roman' -> (9.5, 'Times New Roman')."""
    f = props.get("font", "")
    m = re.match(r"([\d.]+)px\s+'?([^';]+)'?", f)
    if not m:
        return 9.5, "Times New Roman"
    return float(m.group(1)), m.group(2).strip().strip("'")


def px(props: dict, key: str) -> float:
    m = re.match(r"([\d.]+)px", props.get(key, ""))
    return float(m.group(1)) if m else 0.0


class Block(HTMLParser):
    """Плоский разбор: последовательность блоков (параграф / таблица)."""

    def __init__(self):
        super().__init__(convert_charrefs=False)
        self.blocks = []
        self.stack = []            # вложенность таблиц (тут всегда 1 уровень)
        self.cur_p = None
        self.cur_cell = None
        self.fmt = {"b": 0, "u": 0, "i": 0, "spread": 0}

    # --- служебное
    def _emit_p(self, par):
        if self.cur_cell is not None:
            self.cur_cell["paras"].append(par)
        else:
            self.blocks.append({"kind": "p", **par})

    def handle_starttag(self, tag, attrs):
        a = dict(attrs)
        cls = a.get("class", "")
        if tag == "table":
            tbl = {"kind": "table", "rows": []}
            self.blocks.append(tbl)
            self.stack.append(tbl)
        elif tag == "tr" and self.stack:
            self.stack[-1]["rows"].append([])
        elif tag == "td" and self.stack:
            self.cur_cell = {
                "cls": cls, "paras": [],
                "rowspan": int(a.get("rowspan", 1)),
                "colspan": int(a.get("colspan", 1)),
                "valign": a.get("valign", "top"),
            }
            self.stack[-1]["rows"][-1].append(self.cur_cell)
        elif tag == "p":
            self.cur_p = {"cls": cls, "runs": []}
        elif tag == "b":
            self.fmt["b"] += 1
        elif tag == "i":
            self.fmt["i"] += 1
        elif tag == "span":
            if cls == "s1":
                self.fmt["u"] += 1
            elif cls == "s2":
                self.fmt["spread"] += 1
        elif tag == "br" and self.cur_p is not None:
            self.cur_p["runs"].append({"text": "\n", **{k: bool(v) for k, v in self.fmt.items()}})

    def handle_endtag(self, tag):
        if tag == "table" and self.stack:
            self.stack.pop()
        elif tag == "td":
            self.cur_cell = None
        elif tag == "p" and self.cur_p is not None:
            self._emit_p(self.cur_p)
            self.cur_p = None
        elif tag in ("b", "i"):
            self.fmt[tag] = max(0, self.fmt[tag] - 1)
        elif tag == "span":
            if self.fmt["u"]:
                self.fmt["u"] -= 1
            elif self.fmt["spread"]:
                self.fmt["spread"] -= 1

    def handle_data(self, data):
        if self.cur_p is None or not data:
            return
        self.cur_p["runs"].append({"text": data, **{k: bool(v) for k, v in self.fmt.items()}})

    def handle_entityref(self, name):
        self.handle_data(htmllib.unescape(f"&{name};"))

    def handle_charref(self, name):
        self.handle_data(htmllib.unescape(f"&#{name};"))


# --- запись в docx -----------------------------------------------------------

def set_borders(cell, props: dict):
    """Границы ячейки из CSS border-width/border-color (порядок: top right bottom left)."""
    widths = props.get("border-width", "0 0 0 0").replace("px", "").split()
    colors = props.get("border-color", "transparent " * 4).split()
    widths = (widths + ["0"] * 4)[:4]
    colors = (colors + ["transparent"] * 4)[:4]
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for side, w, c in zip(("top", "right", "bottom", "left"), widths, colors):
        el = OxmlElement(f"w:{side}")
        if float(w or 0) > 0 and c != "transparent":
            el.set(qn("w:val"), "single")
            el.set(qn("w:sz"), "4")          # 0.5 pt
            el.set(qn("w:color"), c.lstrip("#").upper())
        else:
            el.set(qn("w:val"), "nil")
        borders.append(el)
    tcPr.append(borders)


def set_fixed_layout(table, widths_mm):
    """Ширины колонок держатся только при fixed layout + явном tblGrid.

    Без этого Word и QuickLook пересчитывают таблицу по содержимому, и колонка
    с длинным текстом схлопывается.
    """
    tblPr = table._tbl.tblPr
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tblPr.append(layout)
    tbl_w = OxmlElement("w:tblW")
    tbl_w.set(qn("w:w"), str(int(sum(widths_mm) * 56.6929)))   # мм -> twips
    tbl_w.set(qn("w:type"), "dxa")
    tblPr.append(tbl_w)
    grid = table._tbl.find(qn("w:tblGrid"))
    if grid is not None:
        table._tbl.remove(grid)
    margins = OxmlElement("w:tblCellMar")
    for side in ("left", "right"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:w"), "57")               # 0.1 см вместо дефолтных 0.19 см
        el.set(qn("w:type"), "dxa")
        margins.append(el)
    tblPr.append(margins)
    grid = OxmlElement("w:tblGrid")
    for mm in widths_mm:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(int(mm * 56.6929)))
        grid.append(col)
    table._tbl.insert(list(table._tbl).index(tblPr) + 1, grid)
    for row in table.rows:
        for ci, cell in enumerate(row.cells):
            if ci < len(widths_mm):
                cell.width = Mm(widths_mm[ci])


def set_valign(cell, valign: str):
    tcPr = cell._tc.get_or_add_tcPr()
    el = OxmlElement("w:vAlign")
    el.set(qn("w:val"), {"bottom": "bottom", "middle": "center"}.get(valign, "top"))
    tcPr.append(el)


def fill_paragraph(par, pdata, p_styles, in_table=False):
    props = p_styles.get(pdata["cls"], {})
    size, family = font_of(props)
    family = FONT_FAMILY or family
    if in_table:
        size = TABLE_SIZE_MAP.get(size, size)
    par.alignment = ALIGN.get(props.get("text-align", "left"))
    pf = par.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(float(props.get("margin", "0 0 0 0").split()[2].replace("px", "") or 0))
    indent = px(props, "text-indent")
    if indent:
        pf.first_line_indent = Pt(indent)
    lh = px(props, "line-height")
    if lh:
        pf.line_spacing = Pt(lh)
    text_written = False
    for run_data in pdata["runs"]:
        chunks = run_data["text"].split("\n")
        for i, chunk in enumerate(chunks):
            if i:
                par.add_run().add_break()
            if not chunk:
                continue
            r = par.add_run(chunk)
            r.font.name = family
            r.font.size = Pt(size)
            r.font.bold = run_data.get("b") or None
            r.font.italic = run_data.get("i") or None
            r.font.underline = run_data.get("u") or None
            r.font.color.rgb = RGBColor(0, 0, 0)
            rpr = r._element.get_or_add_rPr()
            rf = OxmlElement("w:rFonts")
            for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
                rf.set(qn(attr), family)
            rpr.append(rf)
            if run_data.get("spread"):
                sp = OxmlElement("w:spacing")
                sp.set(qn("w:val"), "12")     # 0.6 pt разрядки = 12 twips
                rpr.append(sp)
            text_written = True
    if not text_written and not pdata["runs"]:
        r = par.add_run("")
        r.font.name = family
        r.font.size = Pt(size)


def cell_width(cell_data, td_styles) -> float:
    return px(td_styles.get(cell_data["cls"], {}), "min-width")


def merge_hairlines(rows, td_styles, limit=12.0):
    """Пустая ячейка тоньше 12 pt — не колонка, а полоска: прирастает к соседу.

    Cocoa пишет их как отдельные <td>, из-за чего общая сетка таблицы получает
    колонку в 0.4 мм и весь ряд уезжает.
    """
    for row in rows:
        i = 0
        while i < len(row) - 1:
            w = cell_width(row[i], td_styles)
            empty = not any(r["text"].strip() for p in row[i]["paras"] for r in p["runs"])
            if w < limit and empty:
                nxt = row[i + 1]
                nxt["_extra"] = nxt.get("_extra", 0.0) + w
                nxt["rowspan"] = max(nxt["rowspan"], row[i]["rowspan"])
                del row[i]
                continue
            i += 1
    return rows


def layout(rows, td_styles):
    """Раскладка ячеек по горизонтали с учетом rowspan-ов сверху.

    Cocoa не пишет ни colspan, ни пропуски: строка под вертикально объединенной
    ячейкой содержит только «свои» td, и наивная сумма ширин начинает ряд с нуля.
    Поэтому занятые сверху интервалы держим явно и сдвигаем курсор через них.
    """
    busy = {}                       # индекс строки -> [(x0, x1), ...]
    placed = []                     # [(row_index, cell, x0, width), ...]
    for ri, row in enumerate(rows):
        x = 0.0
        for cell_data in row:
            w = cell_width(cell_data, td_styles) + cell_data.get("_extra", 0.0)
            moved = True
            while moved:
                moved = False
                for x0, x1 in busy.get(ri, []):
                    if x0 - 0.6 <= x < x1 - 0.6:
                        x = x1
                        moved = True
            placed.append((ri, cell_data, x, w))
            if cell_data["rowspan"] > 1:
                for rr in range(ri + 1, min(ri + cell_data["rowspan"], len(rows))):
                    busy.setdefault(rr, []).append((x, x + w))
            x += w
    return placed


def build_grid(placed, tol=2.0):
    """Сетка колонок по границам всех ячеек.

    Границы, разошедшиеся на пару пунктов, — след того, что Cocoa считает ширину
    с паддингами по-разному в разных строках; их надо слить, иначе таблица
    получает колонку в доли миллиметра. Но зазор, равный ширине реальной ячейки
    (колонка «№» в Заявке — 15.6 pt), сливать нельзя: колонка исчезнет. Поэтому
    слияние запрещено, когда разница совпадает с шириной какой-нибудь ячейки.
    """
    weight = {}
    for _, _, x, w in placed:
        for edge in (round(x, 1), round(x + w, 1)):
            weight[edge] = weight.get(edge, 0) + 1
    real_widths = sorted({round(w, 1) for _, _, _, w in placed if w > 0})

    def is_real_column(delta):
        return any(abs(delta - w) <= 2.0 for w in real_widths)

    merged = []
    for edge in sorted(weight):
        if merged:
            delta = edge - merged[-1]
            if delta < tol and not is_real_column(delta):
                if weight[edge] > weight[merged[-1]]:
                    merged[-1] = edge          # оставляем ту границу, что встречается чаще
                continue
        merged.append(edge)
    return merged


def nearest(edges, value):
    return min(range(len(edges)), key=lambda i: abs(edges[i] - value))


def backup_if_edited(out_path: str) -> None:
    """Снимок файла перед перезаписью, если его правил кто-то кроме нас.

    Прецедент 20.08.2026: владелец правил выданный .docx в Word, следующая
    пересборка молча затерла правки, а восстанавливать было нечем — Time
    Machine не настроен, автокопии Word нет. Снимок кладется рядом, в
    `.agent/drafts/_baselines/`, с меткой времени в имени: перезаписи там нет,
    значит и хук на неприкосновенность базы «ДО» не нарушается.
    """
    if not os.path.exists(out_path) or "/cases/" not in out_path.replace(os.sep, "/"):
        return
    case_dir = out_path
    while case_dir != "/" and not os.path.isdir(os.path.join(case_dir, ".agent")):
        case_dir = os.path.dirname(case_dir)
    if case_dir == "/":
        return
    dst_dir = os.path.join(case_dir, ".agent", "drafts", "_baselines")
    os.makedirs(dst_dir, exist_ok=True)
    stamp = datetime.datetime.fromtimestamp(os.path.getmtime(out_path)).strftime("%Y%m%d-%H%M%S")
    base, ext = os.path.splitext(os.path.basename(out_path))
    dst = os.path.join(dst_dir, f"{base}-{stamp}{ext}")
    if not os.path.exists(dst):
        shutil.copy2(out_path, dst)
        print(f"снимок прежней версии: {dst}")


def build(blocks, p_styles, td_styles, out_path):
    doc = Document()
    sec = doc.sections[0]
    sec.page_width, sec.page_height = Mm(PAGE_W), Mm(PAGE_H)
    sec.top_margin, sec.bottom_margin = Mm(M_TOP), Mm(M_BOTTOM)
    sec.left_margin, sec.right_margin = Mm(M_LEFT), Mm(M_RIGHT)
    normal = doc.styles["Normal"]
    normal.font.name = FONT_FAMILY or "Times New Roman"
    normal.font.size = Pt(9.5)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.line_spacing = 1.0

    first_par_used = False
    for bi, block in enumerate(blocks):
        if block["kind"] == "p":
            if not first_par_used:
                par = doc.paragraphs[0] if doc.paragraphs else doc.add_paragraph()
                first_par_used = True
            else:
                par = doc.add_paragraph()
            fill_paragraph(par, block, p_styles)
        else:
            rows = merge_hairlines(block["rows"], td_styles)
            if not rows:
                continue
            if block.get("page_break"):
                doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
            placed = layout(rows, td_styles)
            edges = build_grid(placed)
            ncols = len(edges) - 1
            table = doc.add_table(rows=len(rows), cols=ncols)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.autofit = False
            for ri, cell_data, x, w in placed:
                ci = nearest(edges, x)
                ci_end = max(nearest(edges, x + w), ci + 1)
                span = min(ci_end, ncols) - ci
                cell = table.cell(ri, ci)
                props = td_styles.get(cell_data["cls"], {})
                if span > 1:
                    cell = cell.merge(table.cell(ri, ci + span - 1))
                if cell_data["rowspan"] > 1:
                    last_row = min(ri + cell_data["rowspan"] - 1, len(rows) - 1)
                    cell = cell.merge(table.cell(last_row, ci + span - 1))
                set_borders(cell, props)
                set_valign(cell, cell_data["valign"])
                for pi, pdata in enumerate(cell_data["paras"]):
                    par = cell.paragraphs[0] if pi == 0 else cell.add_paragraph()
                    fill_paragraph(par, pdata, p_styles, in_table=True)
            widths_mm = [(edges[ci + 1] - edges[ci]) * PT_TO_MM * SCALE for ci in range(ncols)]
            set_fixed_layout(table, widths_mm)
            if bi + 1 < len(blocks) and blocks[bi + 1]["kind"] == "table":
                doc.add_paragraph()

    add_page_numbers(doc)
    backup_if_edited(out_path)
    doc.save(out_path)


SIGN_LEFT = "Исполнитель: Зарубина Ю.С. _______________"
SIGN_RIGHT = "Заказчик: Балихина Е.В. _______________"


def _footer_signatures(par):
    """Подписи сторон: Исполнитель у левого поля, Заказчик у правого.

    Парафирование каждого листа: подписанный лист нельзя подменить после
    подписания, поэтому подписи идут в колонтитуле, а не только в конце.
    """
    from docx.enum.text import WD_TAB_ALIGNMENT
    pf = par.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.tab_stops.add_tab_stop(Mm(TEXT_W_MM), WD_TAB_ALIGNMENT.RIGHT)
    run = par.add_run(f"{SIGN_LEFT}\t{SIGN_RIGHT}")
    run.font.name = FONT_FAMILY or "Times New Roman"
    run.font.size = Pt(FOOTER_SIZE)


def _footer_page_number(par):
    par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    par.paragraph_format.space_before = Pt(0)
    par.paragraph_format.space_after = Pt(0)
    run = par.add_run()
    run.font.name = FONT_FAMILY or "Times New Roman"
    run.font.size = Pt(FOOTER_SIZE)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def add_page_numbers(doc):
    """Колонтитул: подписи сторон и номер страницы на КАЖДОМ листе, включая первый."""
    sec = doc.sections[0]
    sec.different_first_page_header_footer = False
    footer = sec.footer
    paras = footer.paragraphs
    par = paras[0] if paras else footer.add_paragraph()
    _footer_signatures(par)
    _footer_page_number(footer.add_paragraph())


def block_text(block) -> str:
    if block["kind"] == "p":
        return "".join(r["text"] for r in block["runs"])
    return " ".join(r["text"] for row in block["rows"] for c in row
                    for p in c["paras"] for r in p["runs"])


def mark_appendix_breaks(blocks):
    """Приложение начинается с новой страницы. Хвост пустых абзацев перед ним
    в оригинале и был ручным «переводом страницы» — при явном разрыве он даст
    лишний пустой лист, поэтому срезается."""
    out = []
    for block in blocks:
        if block["kind"] == "table" and "ПРИЛОЖЕНИЕ" in block_text(block):
            while out and out[-1]["kind"] == "p" and not block_text(out[-1]).strip():
                out.pop()
            block["page_break"] = True
        out.append(block)
    return out


def html_of(path: str) -> str:
    """Слепок верстки .doc/.rtf: textutil есть в macOS из коробки, Word не нужен."""
    if path.lower().endswith((".html", ".htm")):
        return open(path, encoding="utf-8").read()
    tmp = os.path.join(tempfile.mkdtemp(), "src.html")
    subprocess.run(["textutil", "-convert", "html", "-encoding", "UTF-8",
                    "-output", tmp, path], check=True)
    return open(tmp, encoding="utf-8").read()


def convert(src_path: str, out_path: str, transform=None) -> None:
    """.doc -> .docx с сохранением верстки. transform(blocks) — правка содержимого."""
    src = html_of(src_path)
    p_styles, td_styles = parse_css(src)
    parser = Block()
    parser.feed(src)
    blocks = parser.blocks
    if transform is not None:
        blocks = transform(blocks)
    build(mark_appendix_breaks(blocks), p_styles, td_styles, out_path)


def main():
    if len(sys.argv) < 3:
        print("использование: doc_legacy_to_docx.py ИСХОДНИК.doc РЕЗУЛЬТАТ.docx", file=sys.stderr)
        return 2
    convert(sys.argv[1], sys.argv[2])
    print("saved", sys.argv[2])
    return 0


if __name__ == "__main__":
    sys.exit(main())
