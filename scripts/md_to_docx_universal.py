#!/usr/bin/env python3
"""
Универсальный конвертер markdown -> docx для судебных документов.
Поддерживает: заголовки, параграфы, маркированные списки, таблицы,
жирный текст (**text**).
"""

import re
import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Гарнитура стандарта практики (DOCX_FORMATTING.md §1, решение владельца
# 03.08.2026). Берем из DocBuilder, чтобы не разъехалось при следующей правке.
try:
    from scripts.create_docx import FONT_BODY as FONT, apply_page_numbers
except ImportError:  # запуск из каталога scripts/
    from create_docx import FONT_BODY as FONT, apply_page_numbers


def _set_font(run, size_pt, bold=False):
    run.font.name = FONT
    run.font.size = Pt(size_pt)
    run.font.bold = bold


def _strip_yo(doc):
    for t in doc.element.iter(qn("w:t")):
        if t.text and ("ё" in t.text or "Ё" in t.text):
            t.text = t.text.replace("ё", "е").replace("Ё", "Е")


def _add_table(doc, lines):
    """lines - строки markdown таблицы (без разделителя)."""
    rows = []
    for line in lines:
        if line.strip().startswith("|"):
            cells = [c.strip() for c in line.strip()[1:].split("|")]
            if cells and not cells[-1].strip():
                cells = cells[:-1]
            rows.append(cells)

    if not rows:
        return

    # отфильтровать пустые строки
    rows = [r for r in rows if any(c.strip() for c in r)]
    if not rows:
        return

    n_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=n_cols)

    # убрать границы
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement("w:tblBorders")
    for side in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "none")
        tblBorders.append(el)
    tblPr.append(tblBorders)

    # автоподбор ширины
    tblPr.append(OxmlElement("w:tblLayout"))
    tblPr.find(qn("w:tblLayout")).set(qn("w:type"), "autofit")

    for i, row_data in enumerate(rows):
        row = table.rows[i]
        for j in range(n_cols):
            cell = row.cells[j]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            text = row_data[j] if j < len(row_data) else ""
            # парсим жирный текст
            parts = re.split(r"\*\*(.*?)\*\*", text)
            for k, part in enumerate(parts):
                if not part:
                    continue
                bold = (k % 2 == 1)
                _set_font(p.add_run(part), 11, bold=bold)


def _parse_bold_text(text):
    parts = re.split(r"\*\*(.*?)\*\*", text)
    result = []
    for k, part in enumerate(parts):
        if not part:
            continue
        result.append((part, k % 2 == 1))
    return result


def _refuse_gotovo(out_path):
    """Судебный документ в GOTOVO/ собирается только под вердиктом Кони.

    Точка сборки филируемого документа одна — DocBuilder (create_docx.py) с
    вердиктным гейтом. Этот конвертер — для внутренних пакетов (напр. prep.docx
    в 02_hearings/). Прямая запись .docx в GOTOVO мимо вердикта запрещена: пока
    рядом стоит второй вход, вердикт соблюдается добровольно.
    ponytail: граница по имени каталога GOTOVO — зарезервированное место готового
    документа; иные внутренние артефакты (02_hearings, drafts, /tmp) свободны.
    """
    if any(p.upper() == "GOTOVO" for p in Path(out_path).resolve().parts):
        print("СТОП, НЕ СОХРАНЕНО: судебный документ в GOTOVO собирается только "
              "DocBuilder под вердиктом Кони (scripts/create_docx.py). "
              "Этот конвертер — для внутренних пакетов (prep.docx и т.п.).")
        sys.exit(2)


def md_to_docx(md_path, out_path, margins_mm=None, page_numbers=True):
    """margins_mm — (верх, низ, лево, право) в миллиметрах; по умолчанию
    судебный профиль 20/30/30/15 (нижнее поле 30 мм — зона штампа суда).
    Кадровым и договорным документам штамп не нужен: они зовут с (20, 20, 25, 20).

    page_numbers по умолчанию включены: «номера страниц обязательны в каждом
    документе без исключений» (DOCX_FORMATTING.md), а этот вход их не ставил
    вовсе — внутренние пакеты уходили без нумерации.
    """
    md_path = Path(md_path)
    out_path = Path(out_path)
    _refuse_gotovo(out_path)

    text = md_path.read_text(encoding="utf-8")
    lines = text.splitlines()

    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    verh, niz, levo, pravo = margins_mm or (20, 30, 30, 15)
    section.top_margin = Cm(verh / 10)
    section.bottom_margin = Cm(niz / 10)   # 30 мм: зона штампа суда, DOCX_FORMATTING.md
    section.left_margin = Cm(levo / 10)
    section.right_margin = Cm(pravo / 10)

    # удалить дефолтный пустой параграф
    for p in doc.paragraphs:
        p._element.getparent().remove(p._element)

    i = 0
    header_lines = []
    in_header = True

    while i < len(lines):
        line = lines[i]

        # разделитель шапки
        if in_header and line.strip() == "---":
            in_header = False
            i += 1
            continue

        # заголовок первого уровня - основной заголовок документа
        if re.match(r"^# ", line):
            in_header = False
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(4)
            _set_font(p.add_run(line[2:].strip().upper()), 14, bold=True)
            i += 1
            continue

        # подзаголовок второго уровня
        if re.match(r"^## ", line):
            in_header = False
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(4)
            _set_font(p.add_run(line[3:].strip()), 13, bold=True)
            i += 1
            continue

        # раздел
        if re.match(r"^### ", line):
            in_header = False
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
            _set_font(p.add_run(line[4:].strip()), 12, bold=True)
            i += 1
            continue

        # таблица
        if line.strip().startswith("|"):
            in_header = False
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i])
                i += 1
            # пропустить разделительную строку |---|---|
            table_lines = [
                l for l in table_lines
                if not re.match(r"^\|[-:\s|]+\|$", l.strip())
            ]
            _add_table(doc, table_lines)
            continue

        # маркированный список
        if line.strip().startswith("- "):
            in_header = False
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.left_indent = Cm(1.5)
            item_text = line.strip()[2:]
            for part, bold in _parse_bold_text(item_text):
                _set_font(p.add_run(part), 12, bold=bold)
            i += 1
            continue

        # пустая строка
        if not line.strip():
            if in_header:
                pass  # просто пропускаем в шапке
            i += 1
            continue

        # обычный параграф
        in_header = False
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_after = Pt(6)
        if header_lines:
            # это все еще может быть шапка - левое выравнивание
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT

        for part, bold in _parse_bold_text(line.strip()):
            _set_font(p.add_run(part), 12, bold=bold)
        i += 1

    if page_numbers:
        apply_page_numbers(doc)
    _strip_yo(doc)
    doc.save(out_path)
    print(f"Создано: {out_path}")


if __name__ == "__main__":
    if len(sys.argv) < 3:
        print("Использование: python scripts/md_to_docx_universal.py input.md output.docx")
        sys.exit(1)
    md_to_docx(sys.argv[1], sys.argv[2])
