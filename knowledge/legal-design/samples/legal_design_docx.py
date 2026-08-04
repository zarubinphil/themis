#!/usr/bin/env python3
"""Приемы legal design поверх DocBuilder — то, чего в нем пока нет.

Все функции проверены рендером (QuickLook macOS) и разбором XML.
Ключевой урок зонда 03.08.2026: OOXML требует СТРОГОГО порядка дочерних
элементов внутри w:pPr и w:tcPr. Элемент, приписанный в конец, схема
считает недопустимым, и рендерер молча игнорирует его целиком — рамка
абзаца и заливка ячейки исчезают без единой ошибки.
"""
from pathlib import Path

from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Порядок дочерних элементов по ECMA-376. Обрезан до используемого нами.
PPR_ORDER = [
    "pStyle", "keepNext", "keepLines", "pageBreakBefore", "framePr",
    "widowControl", "numPr", "suppressLineNumbers", "pBdr", "shd", "tabs",
    "suppressAutoHyphens", "kinsoku", "wordWrap", "overflowPunct",
    "topLinePunct", "autoSpaceDE", "autoSpaceDN", "bidi", "adjustRightInd",
    "snapToGrid", "spacing", "ind", "contextualSpacing", "mirrorIndents",
    "suppressOverlap", "jc", "textDirection", "textAlignment",
    "textboxTightWrap", "outlineLvl", "divId", "cnfStyle", "rPr", "sectPr",
]
TCPR_ORDER = [
    "cnfStyle", "tcW", "gridSpan", "hMerge", "vMerge", "tcBorders", "shd",
    "noWrap", "tcMar", "textDirection", "tcFitText", "vAlign", "hideMark",
]

from scripts.create_docx import FONT_BODY, FONT_DISPLAY, FONT_MONO  # noqa: E402

FONT = FONT_BODY


def _ordered_insert(parent, element, order):
    """Вставить element в parent на позицию, требуемую схемой OOXML."""
    tag = element.tag.split("}")[-1]
    idx = order.index(tag)
    for child in parent:
        child_tag = child.tag.split("}")[-1]
        if child_tag not in order or order.index(child_tag) > idx:
            child.addprevious(element)
            return
    parent.append(element)


def _rpr(size_pt=12, bold=False, color=None, underline=False, family=None):
    """Готовый w:rPr для ручной сборки run в поле или гиперссылке."""
    rPr = OxmlElement("w:rPr")
    rf = OxmlElement("w:rFonts")
    for attr in ("w:ascii", "w:hAnsi", "w:cs"):
        rf.set(qn(attr), family or FONT_BODY)
    rPr.append(rf)
    if bold:
        rPr.append(OxmlElement("w:b"))
    # Подчеркивание и цвет задаются ЯВНО всегда: элемент w:hyperlink без них
    # рендерер оформляет стилем Hyperlink — синим с подчеркиванием, а это
    # нарушает эталон (курсива и подчеркивания нет, цвет авто).
    # Порядок внутри w:rPr тоже задан схемой: color → sz → u. Переставишь —
    # элемент невалиден и молча игнорируется, как было с pBdr.
    c = OxmlElement("w:color")
    c.set(qn("w:val"), color or "000000")
    rPr.append(c)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(int(size_pt * 2)))
    rPr.append(sz)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single" if underline else "none")
    rPr.append(u)
    return rPr


def _set_font(run, size_pt, bold=False, family=None):
    run.font.name = family or FONT_BODY
    run.font.size = Pt(size_pt)
    run.font.bold = bold


def _norm_parts(parts):
    """Привести к list[(text, bold)]: строка, список строк и список кортежей
    смешиваются в вызовах постоянно — падать на этом незачем."""
    if isinstance(parts, str):
        return [(parts, False)]
    out = []
    for item in parts:
        out.append((item, False) if isinstance(item, str) else tuple(item))
    return out


# ---------------------------------------------------------------- рамка/фон
# Рамка абзаца (w:pBdr) намеренно НЕ реализована. XML собирается схемно
# верным — все четыре стороны на месте, — но рендерер macOS показал только
# нижнюю границу и проглотил заливку (зонд 03.08.2026). Приема, который
# выглядит по-разному в разных программах, в судебном документе быть не
# должно: рамку дает add_summary_box на таблице 1x1, она отрисовалась везде.


def shade_cell(cell, hex_fill):
    """Заливка ячейки таблицы. Только серые тона: документ печатают ч/б."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_fill)
    _ordered_insert(tcPr, shd, TCPR_ORDER)
    return cell


def shade_header_row(table, hex_fill="E6E6E6"):
    for c in table.rows[0].cells:
        shade_cell(c, hex_fill)
    return table


# ------------------------------------------------------- закладки и ссылки
_BM_ID = [1000]


def add_bookmark(paragraph, name):
    """Якорь для внутренней ссылки. Имя без пробелов, латиницей."""
    _BM_ID[0] += 1
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(_BM_ID[0]))
    start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(_BM_ID[0]))
    paragraph._p.insert(0, start)
    paragraph._p.append(end)
    return paragraph


def append_internal_link(paragraph, anchor, text, size_pt=12, bold=False):
    """Внутренняя ссылка в существующий абзац. Черная, без подчеркивания:
    подчеркивание запрещено эталоном, а синий цвет в ч/б печати сереет."""
    hl = OxmlElement("w:hyperlink")
    hl.set(qn("w:anchor"), anchor)
    r = OxmlElement("w:r")
    r.append(_rpr(size_pt, bold=bold))
    t = OxmlElement("w:t")
    t.text = text
    t.set(qn("xml:space"), "preserve")
    r.append(t)
    hl.append(r)
    paragraph._p.append(hl)
    return paragraph


# --------------------------------------------------------------- навигация
def add_static_toc(b, entries, title="СОДЕРЖАНИЕ", linked=False):
    """Статическое оглавление; при linked=True — со ссылками на закладки.

    Поле TOC не годится: оно вычисляется только Word и только по команде
    обновления полей — в любом стороннем просмотрщике и при первом открытии
    остается пустой строкой (проверено зондом 03.08.2026). Статический
    список живет везде и печатается всегда.

    linked по умолчанию ВЫКЛЮЧЕН. Элемент w:hyperlink рендерер оформляет
    стилем Hyperlink — синим с подчеркиванием — поверх прямого
    форматирования run: явные `w:color 000000` и `w:u none` в схемно верном
    порядке проверены и не помогли (замер 03.08.2026, QuickLook macOS).
    Для судебного документа это прямое нарушение эталона (подчеркивания
    нет, цвет авто), поэтому там оглавление идет обычным текстом. Для
    договора ссылки уместны — включать явно.

    entries: list[(anchor, text)] — anchor должен совпасть с add_bookmark.
    """
    p = b.doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    _set_font(p.add_run(title), 12, bold=True, family=FONT_DISPLAY)
    for anchor, text in entries:
        ip = b.doc.add_paragraph()
        ip.alignment = WD_ALIGN_PARAGRAPH.LEFT
        ip.paragraph_format.left_indent = Cm(0.75)
        ip.paragraph_format.first_line_indent = Cm(0)
        ip.paragraph_format.space_before = Pt(0)
        ip.paragraph_format.space_after = Pt(2)
        if linked:
            append_internal_link(ip, anchor, text)
        else:
            _set_font(ip.add_run(text), 12)
    return p


def add_numbered_body(b, num, parts, indent_cm=1.1):
    """Совместимость: номер теперь ведет Word, аргумент num не используется.

    До 03.08.2026 номер писался текстом, и удаление пункта заставляло
    перенумеровывать остальные руками. Теперь нумерация настоящая
    (DocBuilder.add_numbered_body), а сигнатура сохранена, чтобы не трогать
    вызовы в сборщиках образцов.
    """
    return b.add_numbered_body(parts)




def add_summary_box(b, lines, fill="F2F2F2"):
    """Визуальное суммари: рамка с несколькими строками.

    Реализовано таблицей 1x1, а не рамкой абзаца: абзацная рамка на группе
    абзацев в Word схлопывается в общий контур и ломается при разрыве
    страницы. Таблица переносится корректно.
    """
    t = b.doc.add_table(rows=1, cols=1)
    tbl = t._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    borders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "6")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "auto")
        borders.append(el)
    tblPr.append(borders)

    section = b.doc.sections[0]
    avail = int(section.page_width - section.left_margin
                - section.right_margin) // 635
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"), str(avail))
    tblW.set(qn("w:type"), "dxa")
    tblPr.append(tblW)

    cell = t.rows[0].cells[0]
    if fill:
        shade_cell(cell, fill)
    cell.paragraphs[0]._element.getparent().remove(cell.paragraphs[0]._element)
    for parts in lines:
        p = cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
        for text, bold in _norm_parts(parts):
            _set_font(p.add_run(text), 12, bold=bold)
    return t


# ------------------------------------------------------------------ схемы
def timeline_png(path, events, width_in=6.3, height_in=1.9):
    """Таймлайн событий дела в PNG для вставки в .docx.

    events: list[(дата ДД.ММ.ГГГГ, подпись)]. Только ч/б: цвет в судебном
    документе не выживает печать.
    """
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt
    from matplotlib import rcParams
    rcParams["font.family"] = FONT_BODY

    fig, ax = plt.subplots(figsize=(width_in, height_in), dpi=300)
    xs = list(range(len(events)))
    ax.plot(xs, [0] * len(xs), color="#000000", linewidth=0.9, zorder=1)
    ax.scatter(xs, [0] * len(xs), s=38, color="#000000", zorder=2)
    for i, (date, label) in enumerate(events):
        ax.text(i, 0.22, date, ha="center", va="bottom", fontsize=7.6)
        wrapped = "\n".join(_wrap(label, 18))
        ax.text(i, -0.26, wrapped, ha="center", va="top", fontsize=7.6)
    ax.set_ylim(-1.35, 0.85)
    ax.set_xlim(-0.62, len(events) - 0.38)
    ax.axis("off")
    fig.tight_layout(pad=0.15)
    fig.savefig(path, dpi=300, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    return path


def parties_png(path, center, links, width_in=6.3, wrap_at=20):
    """Схема связей сторон: слева сторона-центр, справа столбцом остальные.

    Раскладка «звездой» по кругу первой версии наезжала боксами друг на
    друга — при пяти участниках схема становилась нечитаемой. Столбец
    геометрически не пересекается никогда, и читается сверху вниз, как
    привычный юристу перечень лиц, участвующих в деле.

    links: list[(подпись узла, подпись связи)].
    """
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt
    from matplotlib import rcParams
    from matplotlib.patches import FancyBboxPatch
    rcParams["font.family"] = FONT_BODY

    left_lines = _wrap(center, wrap_at)
    right_blocks = [_wrap(label, wrap_at) for label, _ in links]
    n = len(links)

    row_h = 0.92
    height = max(1.9, n * row_h + 0.35)
    fig, ax = plt.subplots(figsize=(width_in, height), dpi=300)

    char_w = 0.058
    lw = max(1.5, char_w * max(len(s) for s in left_lines) * 1.05)
    rw = max(1.5, char_w * max(len(s) for b in right_blocks for s in b) * 1.05)
    x_left, x_right = 0.0, lw + 2.5 + rw
    ax.set_xlim(-lw - 0.12, x_right + rw + 0.12)

    ys = [(n - 1 - i) * row_h - (n - 1) * row_h / 2 for i in range(n)]
    ax.set_ylim(min(ys) - row_h * 0.62, max(ys) + row_h * 0.62)
    ax.axis("off")

    def node(x, y, lines, half_w, bold=False):
        half_h = 0.115 * len(lines) + 0.09
        ax.add_patch(FancyBboxPatch(
            (x - half_w, y - half_h), 2 * half_w, 2 * half_h,
            boxstyle="round,pad=0.015,rounding_size=0.04",
            linewidth=0.9, edgecolor="#000000", facecolor="#FFFFFF", zorder=3))
        ax.text(x, y, "\n".join(lines), ha="center", va="center", fontsize=7.9,
                fontweight="bold" if bold else "normal", zorder=4,
                linespacing=1.25)

    for (label, rel), lines, y in zip(links, right_blocks, ys):
        ax.annotate("", xy=(x_right - rw - 0.04, y), xytext=(lw + 0.04, 0),
                    arrowprops=dict(arrowstyle="-|>", linewidth=0.85,
                                    color="#000000",
                                    shrinkA=0, shrinkB=0,
                                    connectionstyle="arc3,rad=0.0"), zorder=2)
        mx, my = (lw + x_right - rw) / 2, y / 2
        ax.text(mx, my, rel, ha="center", va="center", fontsize=7.0,
                bbox=dict(boxstyle="square,pad=0.16", facecolor="white",
                          edgecolor="none"), zorder=4)
        node(x_right, y, lines, rw, bold=False)
    node(x_left, 0.0, left_lines, lw, bold=True)

    fig.tight_layout(pad=0.1)
    fig.savefig(path, dpi=300, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    return path


def add_picture_centered(b, png_path, width_cm=16.0, caption=None):
    p = b.doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    p.add_run().add_picture(str(png_path), width=Cm(width_cm))
    if caption:
        c = b.doc.add_paragraph()
        c.alignment = WD_ALIGN_PARAGRAPH.CENTER
        c.paragraph_format.space_after = Pt(8)
        _set_font(c.add_run(caption), 11)
    return p


def set_document_font(b, family):
    """Сменить гарнитуру всего документа.

    Только для НЕсудебных документов (договор, оферта, политика): у
    процессуальных гарнитура задана стандартом — PT Serif, одна на весь
    документ (решение владельца 04.08.2026), и меняться не должна.
    Берем гарнитуру, установленную у всех (Arial, Verdana, Georgia), —
    экзотический шрифт на чужой машине подменится, и верстка поедет.
    """
    for rf in b.doc.element.iter(qn("w:rFonts")):
        for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
            if rf.get(qn(attr)) is not None:
                rf.set(qn(attr), family)
    for p in b.doc.paragraphs:
        for r in p.runs:
            r.font.name = family
    for t in b.doc.tables:
        for row in t.rows:
            for c in row.cells:
                for p in c.paragraphs:
                    for r in p.runs:
                        r.font.name = family
    return b


def add_footer_note(b, text, size_pt=11):
    """Нижний колонтитул LEFT, отдельным абзацем от номера страницы.

    Кегль 11 — минимальный по спецификации: на 9 pt образцы не проходили
    собственный `document_guard.py` и потому не годились как база сравнения
    для проверки форматирования. Правый нижний угол остается свободным под
    штампы суда (DOCX_FORMATTING §2)."""
    footer = b.doc.sections[0].footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _set_font(p.add_run(text), size_pt)
    return p


def _wrap(text, width):
    words, line, out = text.split(), "", []
    for w in words:
        if len(line) + len(w) + 1 > width and line:
            out.append(line)
            line = w
        else:
            line = f"{line} {w}".strip()
    if line:
        out.append(line)
    return out
